from docx import Document
from docx.shared import Inches
from io import BytesIO
import numpy as np
import matplotlib.pyplot as plt
import cv2


def imgToFloat(img):
    if np.issubdtype(img.dtype, np.floating):
        img = img[:, :, :3]
        return img
    else:
        img = img / 255.0
        img = img[:, :, :3]
        return img


def generatePalette(N):
    pallete = np.linspace(0, 1, N).reshape(N, 1)
    return pallete


def colorFit(pixel, palette):
    closestColor = palette[np.argmin(np.linalg.norm(palette - pixel, axis=1))]
    return closestColor


def kwantColorFit(img, palette):
    outImg = img.copy()
    for w in range(img.shape[0]):
        for k in range(img.shape[1]):
            outImg[w, k] = colorFit(img[w, k], palette)
    return outImg


def ditheringRandom(img):
    height, width = img.shape[:2]
    r = np.random.rand(height, width)
    img = (img[:, :, 0] >= r) * 1
    return img


def ditheringOrganized(img, palette):
    M2 = np.array([
        [0, 8, 2, 10],
        [12, 4, 14, 6],
        [3, 11, 1, 9],
        [15, 7, 13, 5]
    ])
    n = 2
    Mpre = (M2 + 1) / (2 * n) ** 2 - 0.5

    height, width = img.shape[:2]
    tiledM = np.tile(Mpre, (height // n + 1, width // n + 1))
    tiledM = tiledM[:height, :width]

    for y in range(height):
        for x in range(width):
            newPixel = img[y, x] + (tiledM[y, x])
            img[y, x] = colorFit(newPixel, palette)
    return img


def ditheringFloydSteinberg(img, palette):
    height, width = img.shape[:2]
    for y in range(height):
        for x in range(width):
            oldPixel = img[y, x].copy()
            newPixel = colorFit(oldPixel, palette)
            img[y, x] = newPixel
            qError = oldPixel - newPixel
            if x + 1 < width:
                img[y, x + 1] = img[y, x + 1] + qError * 7 / 16
            if x - 1 >= 0 and y + 1 < height:
                img[y + 1, x - 1] = img[y + 1, x - 1] + qError * 3 / 16
            if y + 1 < height:
                img[y + 1, x] = img[y + 1, x] + qError * 5 / 16
            if x + 1 < width and y + 1 < height:
                img[y + 1, x + 1] = img[y + 1, x + 1] + qError * 1 / 16
    return img


palette8 = np.array([
        [0.0, 0.0, 0.0],
        [0.0, 0.0, 1.0],
        [0.0, 1.0, 0.0],
        [0.0, 1.0, 1.0],
        [1.0, 0.0, 0.0],
        [1.0, 0.0, 1.0],
        [1.0, 1.0, 0.0],
        [1.0, 1.0, 1.0],
])

palette16 = np.array([
        [0.0, 0.0, 0.0],
        [0.0, 1.0, 1.0],
        [0.0, 0.0, 1.0],
        [1.0, 0.0, 1.0],
        [0.0, 0.5, 0.0],
        [0.5, 0.5, 0.5],
        [0.0, 1.0, 0.0],
        [0.5, 0.0, 0.0],
        [0.0, 0.0, 0.5],
        [0.5, 0.5, 0.0],
        [0.5, 0.0, 0.5],
        [1.0, 0.0, 0.0],
        [0.75, 0.75, 0.75],
        [0.0, 0.5, 0.5],
        [1.0, 1.0, 1.0],
        [1.0, 1.0, 0.0]
])

document = Document()
document.add_heading('Laboratorium 04', 0)

images = ['GS_0001.tif', 'GS_0002.png', 'GS_0003.png']
palettes = [1, 2, 4]

for image in images:
    img = cv2.imread(image)
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    img = imgToFloat(img)
    for palette in palettes:
        imgQuantized = kwantColorFit(img.copy(), generatePalette(2 ** palette))
        imgDitherRandom = ditheringRandom(img.copy())
        imgDitherOrganized = ditheringOrganized(img.copy(), generatePalette(2 ** palette))
        imgDitherFloydSteinberg = ditheringFloydSteinberg(img.copy(), generatePalette(2 ** palette))
        memfile = BytesIO()

        if palette == 1:
            fig, axs = plt.subplots(2, 3)
            fig.tight_layout()

            axs[0, 0].imshow(img)
            axs[0, 0].set_title('Original')
            axs[0, 0].axis('off')
            axs[0, 1].imshow(imgQuantized)
            axs[0, 1].set_title('Quantization')
            axs[0, 1].axis('off')
            axs[0, 2].imshow(imgDitherOrganized)
            axs[0, 2].set_title('Organized')
            axs[0, 2].axis('off')
            axs[1, 0].axis('off')
            axs[1, 1].imshow(imgDitherRandom, cmap=plt.cm.gray)
            axs[1, 1].set_title('Random')
            axs[1, 1].axis('off')
            axs[1, 2].imshow(imgDitherFloydSteinberg)
            axs[1, 2].set_title('Floyd-Steinberg')
            axs[1, 2].axis('off')
        else:
            fig, axs = plt.subplots(2, 2)
            fig.tight_layout()

            axs[0, 0].imshow(img)
            axs[0, 0].set_title('Original')
            axs[0, 0].axis('off')
            axs[0, 1].imshow(imgDitherOrganized)
            axs[0, 1].set_title('Organized')
            axs[0, 1].axis('off')
            axs[1, 0].imshow(imgQuantized)
            axs[1, 0].set_title('Quantization')
            axs[1, 0].axis('off')
            axs[1, 1].imshow(imgDitherFloydSteinberg)
            axs[1, 1].set_title('Floyd-Steinberg')
            axs[1, 1].axis('off')

        plt.savefig(memfile, bbox_inches='tight')
        document.add_heading('{} - dithering {}-bit'.format(image, palette), 2)
        document.add_picture(memfile, width=Inches(4.3))
        plt.close(fig)

images = ['SMALL_0001.tif', 'SMALL_0004.jpg', 'SMALL_0006.jpg', 'SMALL_0009.jpg']
palettes = [palette8, palette16]

for image in images:
    img = cv2.imread(image)
    img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
    img = imgToFloat(img)
    for palette in palettes:
        if len(palette) == len(palette8):
            paletteName = '8 colors'
        else:
            paletteName = '16 colors'

        imgQuantized = kwantColorFit(img.copy(), palette)
        imgDitherOrganized = ditheringOrganized(img.copy(), palette)
        imgDitherFloydSteinberg = ditheringFloydSteinberg(img.copy(), palette)

        memfile = BytesIO()
        fig, axs = plt.subplots(2, 2)
        fig.tight_layout()

        axs[0, 0].imshow(img)
        axs[0, 0].set_title('Original')
        axs[0, 0].axis('off')

        axs[0, 1].imshow(imgDitherOrganized)
        axs[0, 1].set_title('Organized')
        axs[0, 1].axis('off')

        axs[1, 0].imshow(imgQuantized)
        axs[1, 0].set_title('Quantization')
        axs[1, 0].axis('off')

        axs[1, 1].imshow(imgDitherFloydSteinberg)
        axs[1, 1].set_title('Floyd-Steinberg')
        axs[1, 1].axis('off')

        plt.savefig(memfile, bbox_inches='tight')
        document.add_heading('{} - dithering {}'.format(image, paletteName), 2)
        document.add_picture(memfile, width=Inches(4.3))
        plt.close(fig)

document.save('lab04.docx')
