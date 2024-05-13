from docx import Document
from docx.shared import Inches
from io import BytesIO
import cv2
import numpy as np
import matplotlib.pyplot as plt

kat = '.\\'  # katalog z plikami wideo
# plik = "clip_4.mp4"  # nazwa pliku
ile = 15  # ile klatek odtworzyć? <0 - całość
key_frame_counter = 6  # co która klatka ma być kluczowa i nie podlegać kompresji
plot_frames = np.array([])  # automatycznie wyrysuj wykresy
auto_pause_frames = np.array([5])  # automatycznie za pauzuj dla klatki
# subsampling = "4:1:1"  # parametry dla chroma subsampling
dzielnik = 4  # dzielnik przy zapisie różnicy
wyswietlaj_kaltki = False  # czy program ma wyświetlać klatki
ROI = [[300, 400, 500, 600]]   # wyświetlane fragmenty (można podać kilka )
# use_key_compress = 0
# use_rle = 0


class data:
    def init(self):
        self.Y = None
        self.Cb = None
        self.Cr = None


def Chroma_subsampling(L, subsampling):
    if subsampling == "4:4:4":
        return L
    elif subsampling == "4:2:2":
        return L[:, ::2]
    elif subsampling == "4:4:0" or subsampling == "4:2:0":
        return L[::2, ::2]
    elif subsampling == "4:1:1":
        return L[:, ::4]
    elif subsampling == "4:1:0":
        return L[::4, ::4]
    else:
        raise ValueError("Unsupported subsampling format: {}".format(subsampling))


def Chroma_resampling(L, subsampling):
    if subsampling == "4:4:4":
        return L
    elif subsampling == "4:2:2":
        resampled_L = np.zeros((L.shape[0], L.shape[1] * 2), dtype=L.dtype)
        resampled_L[:, ::2] = L
        resampled_L[:, 1::2] = L
        return resampled_L
    elif subsampling == "4:4:0" or subsampling == "4:2:0":
        resampled_L = np.zeros((L.shape[0] * 2, L.shape[1] * 2), dtype=L.dtype)
        resampled_L[::2, ::2] = L
        resampled_L[1::2, ::2] = L
        resampled_L[:, 1::2] = resampled_L[:, ::2]
        return resampled_L
    elif subsampling == "4:1:1":
        resampled_L = np.zeros((L.shape[0], L.shape[1] * 4), dtype=L.dtype)
        resampled_L[:, ::4] = L
        resampled_L[:, 1::4] = L
        resampled_L[:, 2::4] = L
        resampled_L[:, 3::4] = L
        return resampled_L
    elif subsampling == "4:1:0":
        resampled_L = np.zeros((L.shape[0] * 4, L.shape[1] * 4), dtype=L.dtype)
        resampled_L[::4, ::4] = L
        resampled_L[1::4, ::4] = L
        resampled_L[2::4, ::4] = L
        resampled_L[3::4, ::4] = L
        resampled_L[:, 1::4] = resampled_L[:, ::4]
        resampled_L[:, 2::4] = resampled_L[:, ::4]
        resampled_L[:, 3::4] = resampled_L[:, ::4]
        return resampled_L
    else:
        raise ValueError("Unsupported subsampling format: {}".format(subsampling))


def rle(img):
    count = 1
    shape = img.shape
    img = np.array(img).flatten()
    compressed_image = list()
    for i in range(len(img)):
        if i+1 == len(img):
            compressed_image.append(count)
            compressed_image.append(img[i])
            break
        if img[i] == img[i+1]:
            count += 1
        else:
            compressed_image.append(count)
            compressed_image.append(img[i])
            count = 1
    return np.array(compressed_image), shape


def rle_decompress(compressed_image):
    compressed_data, shape = compressed_image
    decompressed_image = []
    for i in range(0, len(compressed_data), 2):
        count = compressed_data[i]
        pixel_value = compressed_data[i+1]
        decompressed_image.extend([pixel_value] * count)
    return np.array(decompressed_image).reshape(shape)


def frame_image_to_class(frame, subsampling):
    Frame_class = data()
    Frame_class.Y = frame[:, :, 0].astype(int)
    Frame_class.Cb = Chroma_subsampling(frame[:, :, 2].astype(int), subsampling)
    Frame_class.Cr = Chroma_subsampling(frame[:, :, 1].astype(int), subsampling)
    return Frame_class


def frame_layers_to_image(Y, Cr, Cb, subsampling):
    Cb = Chroma_resampling(Cb, subsampling)
    Cr = Chroma_resampling(Cr, subsampling)
    return np.dstack([Y, Cr, Cb]).clip(0, 255).astype(np.uint8)


def compress_KeyFrame(Frame_class, use_rle, use_key_compress):
    KeyFrame = data()
    KeyFrame.Y = Frame_class.Y
    KeyFrame.Cb = Frame_class.Cb
    KeyFrame.Cr = Frame_class.Cr
    if use_rle and use_key_compress:
        KeyFrame.Y = rle(KeyFrame.Y)
        KeyFrame.Cb = rle(KeyFrame.Cb)
        KeyFrame.Cr = rle(KeyFrame.Cr)
    return KeyFrame


def decompress_KeyFrame(KeyFrame, use_rle, use_key_compress):
    if use_rle and use_key_compress:
        KeyFrame.Y = rle_decompress(KeyFrame.Y)
        KeyFrame.Cb = rle_decompress(KeyFrame.Cb)
        KeyFrame.Cr = rle_decompress(KeyFrame.Cr)
    Y = KeyFrame.Y
    Cb = KeyFrame.Cb
    Cr = KeyFrame.Cr
    frame_image = frame_layers_to_image(Y, Cr, Cb, subsampling)
    return frame_image


def compress_not_KeyFrame(Frame_class, KeyFrame, use_rle, frame_div=dzielnik):
    Compress_data = data()
    Compress_data.Y = (Frame_class.Y - KeyFrame.Y) // frame_div
    Compress_data.Cb = (Frame_class.Cb - KeyFrame.Cb) // frame_div
    Compress_data.Cr = (Frame_class.Cr - KeyFrame.Cr) // frame_div
    if use_rle:
        Compress_data.Y = rle(Compress_data.Y)
        Compress_data.Cb = rle(Compress_data.Cb)
        Compress_data.Cr = rle(Compress_data.Cr)
    return Compress_data


def decompress_not_KeyFrame(Compress_data, KeyFrame, use_rle, frame_div=dzielnik):
    if use_rle:
        Compress_data.Y = rle_decompress(Compress_data.Y)
        Compress_data.Cb = rle_decompress(Compress_data.Cb)
        Compress_data.Cr = rle_decompress(Compress_data.Cr)
    Y = Compress_data.Y * frame_div + KeyFrame.Y
    Cb = Compress_data.Cb * frame_div + KeyFrame.Cb
    Cr = Compress_data.Cr * frame_div + KeyFrame.Cr
    return frame_layers_to_image(Y, Cr, Cb, subsampling)


def plotDiffrence(ReferenceFrame, DecompressedFrame, ROI, document, subsampling, dzielnik_klatki = dzielnik):
    # bardzo słaby i sztuczny przykład wykorzystania tej opcji
    # przerobić żeby porównanie było dokonywane w RGB nie YCrCb i/lub zastąpić innym porównaniem
    # ROI - Region of Insert współrzędne fragmentu który chcemy przybliżyć i ocenić w formacie [w1,w2,k1,k2]
    fig, axs = plt.subplots(4, 3, sharey=True)
    fig.set_size_inches(7, 5)

    #set fig title

    fig.suptitle("Subsampling={}, Divider={}".format( subsampling, dzielnik_klatki))

    ReferenceFrame = cv2.cvtColor(ReferenceFrame, cv2.COLOR_YCrCb2RGB)
    DecompressedFrame = cv2.cvtColor(DecompressedFrame, cv2.COLOR_YCrCb2RGB)

    axs[0][0].imshow(ReferenceFrame[ROI[0]:ROI[1], ROI[2]:ROI[3]])
    axs[0][2].imshow(DecompressedFrame[ROI[0]:ROI[1], ROI[2]:ROI[3]])
    diff = ReferenceFrame[ROI[0]:ROI[1], ROI[2]:ROI[3]].astype(float) - DecompressedFrame[ROI[0]:ROI[1],
                                                                        ROI[2]:ROI[3]].astype(float)
    print(np.min(diff), np.max(diff))
    axs[0][1].imshow(diff, vmin=np.min(diff), vmax=np.max(diff))


    # Display Y, Cb, Cr layers and their differences
    for i in range(3):
        axs[i+1, 0].imshow(ReferenceFrame[:,:,i][ROI[0]:ROI[1], ROI[2]:ROI[3]], cmap=plt.cm.gray)
        axs[i+1, 2].imshow(DecompressedFrame[:,:,i][ROI[0]:ROI[1], ROI[2]:ROI[3]], cmap=plt.cm.gray)
        diff = ReferenceFrame[:,:,i][ROI[0]:ROI[1], ROI[2]:ROI[3]].astype(float) - DecompressedFrame[:,:,i][ROI[0]:ROI[1], ROI[2]:ROI[3]].astype(float)
        axs[i+1, 1].imshow(diff, vmin=np.min(diff), vmax=np.max(diff))


document = Document()
document.add_heading('Laboratorium 09', 0)

pliki = ["clip_3.mp4", "clip_4.mp4", "clip_5.mp4"]
subsamples = ["4:4:4", "4:2:2", "4:4:0", "4:2:0", "4:1:1", "4:1:0"]
methodsKeyCompress = [0, 1]
methodsRLE = [0, 1]

for plik in pliki:
    for methodKeyCompress in methodsKeyCompress:
        if methodKeyCompress == 1:
            use_key_compress = 1
            metodaKeyCompress = "z key compress"
        else:
            use_key_compress = 0
            metodaKeyCompress = "bez key compress"

        for methodRLE in methodsRLE:
            if methodRLE == 1:
                use_rle = 1
                metodaRLE = "z RLE"
            else:
                use_rle = 0
                metodaRLE = "bez RLE"

            for subsampling in subsamples:
                cap = cv2.VideoCapture(kat + '\\' + plik)
                if ile < 0:
                    ile = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))

                # cv2.namedWindow('Normal Frame')
                # cv2.namedWindow('Decompressed Frame')
                ret, frame = cap.read()
                compression_information = np.zeros((3, ile))
                for i in range(ile):
                    if wyswietlaj_kaltki:
                        cv2.imshow('Normal Frame', frame)
                    frame = cv2.cvtColor(frame, cv2.COLOR_BGR2YCrCb)
                    Frame_class = frame_image_to_class(frame, subsampling)
                    if (i % key_frame_counter) == 0:  # pobieranie klatek kluczowych
                        KeyFrame = compress_KeyFrame(Frame_class, use_rle, use_key_compress)
                        cY = KeyFrame.Y
                        cCb = KeyFrame.Cb
                        cCr = KeyFrame.Cr
                        Decompresed_Frame = decompress_KeyFrame(KeyFrame, use_rle, use_key_compress)
                    else:  # kompresja
                        Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame, use_rle)
                        cY = Compress_data.Y
                        cCb = Compress_data.Cb
                        cCr = Compress_data.Cr
                        Decompresed_Frame = decompress_not_KeyFrame(Compress_data, KeyFrame, use_rle)

                    compression_information[0, i] = (frame[:, :, 0].size - cY[0].size) / frame[:, :, 0].size
                    compression_information[1, i] = (frame[:, :, 0].size - cCb[0].size) / frame[:, :, 0].size
                    compression_information[2, i] = (frame[:, :, 0].size - cCr[0].size) / frame[:, :, 0].size
                    if wyswietlaj_kaltki:
                        cv2.imshow('Decompressed Frame', cv2.cvtColor(Decompresed_Frame, cv2.COLOR_YCrCb2BGR))

                    if np.any(plot_frames == i):  # rysuj wykresy
                        for r in ROI:
                            plotDiffrence(frame, Decompresed_Frame, r)

                    if np.any(auto_pause_frames == i):
                        cv2.waitKey(-1)  # wait until any key is pressed

                    k = cv2.waitKey(1) & 0xff

                    if k == ord('q'):
                        break
                    elif k == ord('p'):
                        cv2.waitKey(-1)  # wait until any key is pressed

                memfile = BytesIO()
                plt.figure()
                plt.plot(np.arange(0, ile), compression_information[0, :] * 100)
                plt.plot(np.arange(0, ile), compression_information[1, :] * 100)
                plt.plot(np.arange(0, ile), compression_information[2, :] * 100)
                document.add_heading("File:{}, subsampling={}, divider={}, KeyFrame={}, Metoda {}, {}"
                                     .format(plik, subsampling, dzielnik, key_frame_counter, metodaRLE,
                                             metodaKeyCompress), 3)
                plt.xlabel("Frame number")
                plt.ylabel("Compression ratio [%]")
                plt.legend(["Y", "Cb", "Cr"])
                plt.tight_layout(pad=1.5)
                plt.savefig(memfile, bbox_inches='tight', dpi=500)
                document.add_picture(memfile, width=Inches(3.5))
                plt.close()

document.save('lab09.docx')
