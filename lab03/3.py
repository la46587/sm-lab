from docx import Document
from docx.shared import Inches
from io import BytesIO
import numpy as np
import matplotlib.pyplot as plt
import cv2


def imgToUInt8(img):  # konwersja obrazu na uint8
    if np.issubdtype(img.dtype, np.integer) or np.issubdtype(img.dtype, np.unsignedinteger):
        return img
    else:
        img = img * 255.0
        img = img.astype('uint8')
        img = img[:, :, :3]
        return img


def imgResizeNearestNeighours(img, scale):  # metoda najbliższego sąsiada
    x = img.shape[1]
    y = img.shape[0]

    X = int(np.ceil(x * scale))
    Y = int(np.ceil(y * scale))

    imgResized = np.zeros((Y, X, 3), dtype=np.uint8)

    xx = np.linspace(0, x - 1, X)
    yy = np.linspace(0, y - 1, Y)

    for w in range(Y):
        for k in range(X):
            originalY = min(int(np.round(yy[w])), y - 1)
            originalX = min(int(np.round(xx[k])), x - 1)
            imgResized[w, k] = img[originalY, originalX]

    return imgResized


def interpolate(img, x, y):  # interpolacja
    x0 = int(x)
    y0 = int(y)
    x1 = min(x0 + 1, img.shape[1] - 1)
    y1 = min(y0 + 1, img.shape[0] - 1)

    dx = x - x0
    dy = y - y0

    interpolated_value = ((1 - dx) * (1 - dy) * img[y0, x0] + dx * (1 - dy) * img[y0, x1] +
                          (1 - dx) * dy * img[y1, x0] + dx * dy * img[y1, x1])

    return interpolated_value.astype(np.uint8)


def imgResizeInterpolation(img, scale):  # metoda interpolacji dwuliniowej
    x = img.shape[1]
    y = img.shape[0]

    X = int(np.ceil(x * scale))
    Y = int(np.ceil(y * scale))

    imgResized = np.zeros((Y, X, 3), dtype=np.uint8)

    for w in range(Y):
        for k in range(X):
            originalY = min((w + 0.5) / scale - 0.5, y - 1)
            originalX = min((k + 0.5) / scale - 0.5, x - 1)
            imgResized[w, k] = interpolate(img, originalX, originalY)

    return imgResized


def imgShrinkMean(img, scale):  # metoda średniej
    x = img.shape[1]
    y = img.shape[0]

    X = int(np.ceil(x * scale))
    Y = int(np.ceil(y * scale))

    imgShrunk = np.zeros((Y, X, 3), dtype=np.uint8)

    for w in range(Y):
        for k in range(X):
            originalX = min(int(np.ceil(k / scale)), x - 1)
            originalY = min(int(np.ceil(w / scale)), y - 1)

            neighborsR = []
            neighborsG = []
            neighborsB = []

            for i in range(-1, 2):
                for j in range(-1, 2):
                    if 0 <= originalY + i < y and 0 <= originalX + j < x:
                        pxl = img[originalY + i, originalX + j]
                        neighborsR.append(pxl[0])
                        neighborsG.append(pxl[1])
                        neighborsB.append(pxl[2])

            meanPxlR = np.mean(neighborsR)
            meanPxlG = np.mean(neighborsG)
            meanPxlB = np.mean(neighborsB)
            imgShrunk[w, k] = [meanPxlR, meanPxlG, meanPxlB]

    return imgShrunk


def imgShrinkAverage(img, scale):  # metoda średniej ważonej
    x = img.shape[1]
    y = img.shape[0]

    X = int(np.ceil(x * scale))
    Y = int(np.ceil(y * scale))

    imgShrunk = np.zeros((Y, X, 3), dtype=np.uint8)

    weights = np.array([[1, 2, 1],
                        [2, 4, 2],
                        [1, 2, 1]])
    weights = weights / np.sum(weights)

    for w in range(Y):
        for k in range(X):
            originalX = min(int(np.ceil(k / scale)), x - 1)
            originalY = min(int(np.ceil(w / scale)), y - 1)

            weightedSum = np.zeros(3)
            totalWeight = 0

            for i in range(-1, 2):
                for j in range(-1, 2):
                    if 0 <= originalY + i < y and 0 <= originalX + j < x:
                        weight = weights[i + 1, j + 1]
                        totalWeight += weight
                        weightedSum += img[originalY + i, originalX + j] * weight

            averagePixel = weightedSum / totalWeight
            imgShrunk[w, k] = averagePixel.astype(np.uint8)

    return imgShrunk


def imgShrinkMedian(img, scale):  # metoda mediany
    x = img.shape[1]
    y = img.shape[0]

    X = int(np.ceil(x * scale))
    Y = int(np.ceil(y * scale))

    imgShrunk = np.zeros((Y, X, 3), dtype=np.uint8)

    for w in range(Y):
        for k in range(X):
            originalX = min(int(np.ceil(k / scale)), x - 1)
            originalY = min(int(np.ceil(w / scale)), y - 1)

            neighborsR = []
            neighborsG = []
            neighborsB = []

            for i in range(-1, 2):
                for j in range(-1, 2):
                    if 0 <= originalY + i < y and 0 <= originalX + j < x:
                        pxl = img[originalY + i, originalX + j]
                        neighborsR.append(pxl[0])
                        neighborsG.append(pxl[1])
                        neighborsB.append(pxl[2])

            medianPxlR = np.median(neighborsR)
            medianPxlG = np.median(neighborsG)
            medianPxlB = np.median(neighborsB)
            imgShrunk[w, k] = [medianPxlR, medianPxlG, medianPxlB]

    return imgShrunk


document = Document()
document.add_heading('Laboratorium 03', 0)

imagesBig = ['BIG_0002.jpg', 'BIG_MOON.jpg']
scalesBig = [0.5, 0.25, 0.1]

for image in imagesBig:
    img = plt.imread(image)

    if img.shape[2]:
        img = imgToUInt8(img)

    edge = cv2.Canny(img, 100, 200, L2gradient=True)

    for scale in scalesBig:
        imgMean = imgShrinkMean(img, scale)
        imgAverage = imgShrinkAverage(img, scale)
        imgMedian = imgShrinkMedian(img, scale)

        edgeMean = cv2.Canny(imgMean, 100, 200, L2gradient=True)
        edgeAverage = cv2.Canny(imgAverage, 100, 200, L2gradient=True)
        edgeMedian = cv2.Canny(imgMedian, 100, 200, L2gradient=True)

        memfile = BytesIO()
        fig, axs = plt.subplots(2, 2)
        fig.tight_layout()

        axs[0, 0].imshow(img)
        axs[0, 0].set_xlim(0.65 * img.shape[1], 0.75 * img.shape[1])
        axs[0, 0].set_ylim(0.4 * img.shape[0], 0.3 * img.shape[0])
        axs[0, 0].set_title('Original')

        axs[0, 1].imshow(imgMean)
        axs[0, 1].set_xlim(0.65 * imgMean.shape[1], 0.75 * imgMean.shape[1])
        axs[0, 1].set_ylim(0.4 * imgMean.shape[0], 0.3 * imgMean.shape[0])
        axs[0, 1].set_title('Mean x{}'.format(scale))

        axs[1, 0].imshow(imgAverage)
        axs[1, 0].set_xlim(0.65 * imgAverage.shape[1], 0.75 * imgAverage.shape[1])
        axs[1, 0].set_ylim(0.4 * imgAverage.shape[0], 0.3 * imgAverage.shape[0])
        axs[1, 0].set_title('Average x{}'.format(scale))

        axs[1, 1].imshow(imgMedian)
        axs[1, 1].set_xlim(0.65 * imgMedian.shape[1], 0.75 * imgMedian.shape[1])
        axs[1, 1].set_ylim(0.4 * imgMedian.shape[0], 0.3 * imgMedian.shape[0])
        axs[1, 1].set_title('Median x{}'.format(scale))

        plt.savefig(memfile)
        document.add_heading('Obraz {} - fragment 1'.format(image), 2)
        document.add_picture(memfile, width=Inches(5))
        plt.close(fig)

        memfile = BytesIO()
        fig, axs = plt.subplots(2, 2)
        fig.tight_layout()

        axs[0, 0].imshow(edge)
        axs[0, 0].set_xlim(0.65 * img.shape[1], 0.75 * img.shape[1])
        axs[0, 0].set_ylim(0.4 * img.shape[0], 0.3 * img.shape[0])
        axs[0, 0].set_title('Original')

        axs[0, 1].imshow(edgeMean)
        axs[0, 1].set_xlim(0.65 * imgMean.shape[1], 0.75 * imgMean.shape[1])
        axs[0, 1].set_ylim(0.4 * imgMean.shape[0], 0.3 * imgMean.shape[0])
        axs[0, 1].set_title('Mean x{}'.format(scale))

        axs[1, 0].imshow(edgeAverage)
        axs[1, 0].set_xlim(0.65 * imgAverage.shape[1], 0.75 * imgAverage.shape[1])
        axs[1, 0].set_ylim(0.4 * imgAverage.shape[0], 0.3 * imgAverage.shape[0])
        axs[1, 0].set_title('Average x{}'.format(scale))

        axs[1, 1].imshow(edgeMedian)
        axs[1, 1].set_xlim(0.65 * imgMedian.shape[1], 0.75 * imgMedian.shape[1])
        axs[1, 1].set_ylim(0.4 * imgMedian.shape[0], 0.3 * imgMedian.shape[0])
        axs[1, 1].set_title('Median x{}'.format(scale))

        plt.savefig(memfile)
        document.add_heading('Obraz {} - fragment 1 (krawędzie)'.format(image), 2)
        document.add_picture(memfile, width=Inches(5))
        plt.close(fig)

        memfile = BytesIO()
        fig, axs = plt.subplots(2, 2)
        fig.tight_layout()

        axs[0, 0].imshow(img)
        axs[0, 0].set_xlim(0.55 * img.shape[1], 0.65 * img.shape[1])
        axs[0, 0].set_ylim(0.5 * img.shape[0], 0.4 * img.shape[0])
        axs[0, 0].set_title('Original')

        axs[0, 1].imshow(imgMean)
        axs[0, 1].set_xlim(0.55 * imgMean.shape[1], 0.65 * imgMean.shape[1])
        axs[0, 1].set_ylim(0.5 * imgMean.shape[0], 0.4 * imgMean.shape[0])
        axs[0, 1].set_title('Mean x{}'.format(scale))

        axs[1, 0].imshow(imgAverage)
        axs[1, 0].set_xlim(0.55 * imgAverage.shape[1], 0.65 * imgAverage.shape[1])
        axs[1, 0].set_ylim(0.5 * imgAverage.shape[0], 0.4 * imgAverage.shape[0])
        axs[1, 0].set_title('Average x{}'.format(scale))

        axs[1, 1].imshow(imgMedian)
        axs[1, 1].set_xlim(0.55 * imgMedian.shape[1], 0.65 * imgMedian.shape[1])
        axs[1, 1].set_ylim(0.5 * imgMedian.shape[0], 0.4 * imgMedian.shape[0])
        axs[1, 1].set_title('Median x{}'.format(scale))

        plt.savefig(memfile)
        document.add_heading('Obraz {} - fragment 2'.format(image), 2)
        document.add_picture(memfile, width=Inches(5))
        plt.close(fig)

        memfile = BytesIO()
        fig, axs = plt.subplots(2, 2)
        fig.tight_layout()

        axs[0, 0].imshow(edge)
        axs[0, 0].set_xlim(0.55 * img.shape[1], 0.65 * img.shape[1])
        axs[0, 0].set_ylim(0.5 * img.shape[0], 0.4 * img.shape[0])
        axs[0, 0].set_title('Original')

        axs[0, 1].imshow(edgeMean)
        axs[0, 1].set_xlim(0.55 * imgMean.shape[1], 0.65 * imgMean.shape[1])
        axs[0, 1].set_ylim(0.5 * imgMean.shape[0], 0.4 * imgMean.shape[0])
        axs[0, 1].set_title('Mean x{}'.format(scale))

        axs[1, 0].imshow(edgeAverage)
        axs[1, 0].set_xlim(0.55 * imgAverage.shape[1], 0.65 * imgAverage.shape[1])
        axs[1, 0].set_ylim(0.5 * imgAverage.shape[0], 0.4 * imgAverage.shape[0])
        axs[1, 0].set_title('Average x{}'.format(scale))

        axs[1, 1].imshow(edgeMedian)
        axs[1, 1].set_xlim(0.55 * imgMedian.shape[1], 0.65 * imgMedian.shape[1])
        axs[1, 1].set_ylim(0.5 * imgMedian.shape[0], 0.4 * imgMedian.shape[0])
        axs[1, 1].set_title('Median x{}'.format(scale))

        plt.savefig(memfile)
        document.add_heading('Obraz {} - fragment 2 (krawędzie)'.format(image), 2)
        document.add_picture(memfile, width=Inches(5))
        plt.close(fig)

        memfile = BytesIO()
        fig, axs = plt.subplots(2, 2)
        fig.tight_layout()

        axs[0, 0].imshow(img)
        axs[0, 0].set_xlim(0.45 * img.shape[1], 0.55 * img.shape[1])
        axs[0, 0].set_ylim(0.4 * img.shape[0], 0.3 * img.shape[0])
        axs[0, 0].set_title('Original')

        axs[0, 1].imshow(imgMean)
        axs[0, 1].set_xlim(0.45 * imgMean.shape[1], 0.55 * imgMean.shape[1])
        axs[0, 1].set_ylim(0.4 * imgMean.shape[0], 0.3 * imgMean.shape[0])
        axs[0, 1].set_title('Mean x{}'.format(scale))

        axs[1, 0].imshow(imgAverage)
        axs[1, 0].set_xlim(0.45 * imgAverage.shape[1], 0.55 * imgAverage.shape[1])
        axs[1, 0].set_ylim(0.4 * imgAverage.shape[0], 0.3 * imgAverage.shape[0])
        axs[1, 0].set_title('Average x{}'.format(scale))

        axs[1, 1].imshow(imgMedian)
        axs[1, 1].set_xlim(0.45 * imgMedian.shape[1], 0.55 * imgMedian.shape[1])
        axs[1, 1].set_ylim(0.4 * imgMedian.shape[0], 0.3 * imgMedian.shape[0])
        axs[1, 1].set_title('Median x{}'.format(scale))

        plt.savefig(memfile)
        document.add_heading('Obraz {} - fragment 3'.format(image), 2)
        document.add_picture(memfile, width=Inches(5))
        plt.close(fig)

        memfile = BytesIO()
        fig, axs = plt.subplots(2, 2)
        fig.tight_layout()

        axs[0, 0].imshow(edge)
        axs[0, 0].set_xlim(0.45 * img.shape[1], 0.55 * img.shape[1])
        axs[0, 0].set_ylim(0.4 * img.shape[0], 0.3 * img.shape[0])
        axs[0, 0].set_title('Original')

        axs[0, 1].imshow(edgeMean)
        axs[0, 1].set_xlim(0.45 * imgMean.shape[1], 0.55 * imgMean.shape[1])
        axs[0, 1].set_ylim(0.4 * imgMean.shape[0], 0.3 * imgMean.shape[0])
        axs[0, 1].set_title('Mean x{}'.format(scale))

        axs[1, 0].imshow(edgeAverage)
        axs[1, 0].set_xlim(0.45 * imgAverage.shape[1], 0.55 * imgAverage.shape[1])
        axs[1, 0].set_ylim(0.4 * imgAverage.shape[0], 0.3 * imgAverage.shape[0])
        axs[1, 0].set_title('Average x{}'.format(scale))

        axs[1, 1].imshow(edgeMedian)
        axs[1, 1].set_xlim(0.45 * imgMedian.shape[1], 0.55 * imgMedian.shape[1])
        axs[1, 1].set_ylim(0.4 * imgMedian.shape[0], 0.3 * imgMedian.shape[0])
        axs[1, 1].set_title('Median x{}'.format(scale))

        plt.savefig(memfile)
        document.add_heading('Obraz {} - fragment 3 (krawędzie)'.format(image), 2)
        document.add_picture(memfile, width=Inches(5))
        plt.close(fig)

imagesSmall = ['SMALL_0001.tif', 'SMALL_0003.png', 'SMALL_0005.jpg', 'SMALL_CAT.jpg']
scalesSmall = [1.5, 2, 5]

for image in imagesSmall:
    img = plt.imread(image)

    if img.shape[2]:
        img = imgToUInt8(img)

    edge = cv2.Canny(img, 100, 200, L2gradient=True)

    for scale in scalesSmall:
        imgNN = imgResizeNearestNeighours(img, scale)
        imgInterp = imgResizeInterpolation(img, scale)

        edgeNN = cv2.Canny(imgNN, 100, 200, L2gradient=True)
        edgeInterp = cv2.Canny(imgInterp, 100, 200, L2gradient=True)

        memfile = BytesIO()
        fig, axs = plt.subplots(1, 3)
        fig.tight_layout()

        axs[0].imshow(img)
        axs[0].set_title('Original')

        axs[1].imshow(imgNN)
        axs[1].set_title('Nearest Neighbours x{}'.format(scale))

        axs[2].imshow(imgInterp)
        axs[2].set_title('Interpolation x{}'.format(scale))

        plt.savefig(memfile)
        document.add_heading('Obraz {}'.format(image), 2)
        document.add_picture(memfile, width=Inches(5))
        plt.close(fig)

        memfile = BytesIO()
        fig, axs = plt.subplots(1, 3)
        fig.tight_layout()

        axs[0].imshow(edge)
        axs[0].set_title('Original')

        axs[1].imshow(edgeNN)
        axs[1].set_title('Nearest Neighbours x{}'.format(scale))

        axs[2].imshow(edgeInterp)
        axs[2].set_title('Interpolation x{}'.format(scale))

        plt.savefig(memfile)
        document.add_heading('Obraz {} (krawędzie)'.format(image), 2)
        document.add_picture(memfile, width=Inches(5))
        plt.close(fig)

document.save('lab03.docx')
