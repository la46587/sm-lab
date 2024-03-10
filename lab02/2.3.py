from docx import Document
from docx.shared import Inches
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import cv2


def imgToUInt8(img):
    if np.issubdtype(img.dtype, np.integer) or np.issubdtype(img.dtype, np.unsignedinteger):
        return img
    else:
        img = img * 255.0
        img = img.astype('uint8')
        return img


def imgToFloat(img):
    if np.issubdtype(img.dtype, np.floating):
        return img
    else:
        img = img / 255.0
        return img


df = pd.DataFrame(data={
    'Filename': ['B02.jpg'],
    'Fragments': [[[0, 0, 199, 199], [600, 600, 799, 799]]]
})

files = []

for index, row in df.iterrows():
    img = plt.imread(row['Filename'])
    if row['Fragments'] is not None:
        for f in row['Fragments']:
            frag = img[f[0]:f[2], f[1]:f[3]].copy()
            imgRG = frag.copy()
            imgGG = frag.copy()
            imgBG = frag.copy()
            imgR = frag.copy()
            imgG = frag.copy()
            imgB = frag.copy()

            imgRG = imgRG[:, :, 0]
            imgGG = imgGG[:, :, 1]
            imgBG = imgBG[:, :, 2]

            imgY1 = 0.299 * imgRG + 0.587 * imgGG + 0.114 * imgBG
            imgY2 = 0.2126 * imgRG + 0.7152 * imgGG + 0.0722 * imgBG

            imgR[:, :, 1] = 0
            imgR[:, :, 2] = 0
            imgG[:, :, 0] = 0
            imgG[:, :, 2] = 0
            imgB[:, :, 0] = 0
            imgB[:, :, 1] = 0

            fig, axs = plt.subplots(3, 3)
            fig.tight_layout()
            axs[0, 0].imshow(frag)
            axs[0, 0].set_title('oryginał')
            axs[0, 1].imshow(imgY1, cmap=plt.cm.gray)
            axs[0, 1].set_title('Y1')
            axs[0, 2].imshow(imgY2, cmap=plt.cm.gray)
            axs[0, 2].set_title('Y2')
            axs[1, 0].imshow(imgRG, cmap=plt.cm.gray)
            axs[1, 0].set_title('R')
            axs[1, 1].imshow(imgGG, cmap=plt.cm.gray)
            axs[1, 1].set_title('G')
            axs[1, 2].imshow(imgBG, cmap=plt.cm.gray)
            axs[1, 2].set_title('B')
            axs[2, 0].imshow(imgR)
            axs[2, 0].set_title('R', color='red')
            axs[2, 1].imshow(imgG)
            axs[2, 1].set_title('G', color='green')
            axs[2, 2].imshow(imgB)
            axs[2, 2].set_title('B', color='blue')
            plt.savefig("frag{}".format(f))
            files.append("frag{}".format(f) + '.png')
            plt.show()

document = Document()
document.add_heading('Laboratorium 02 - Zadanie 3', 0)

for file in files:
    document.add_heading('Plik - {}'.format(file), 2)

    memfile = BytesIO()
    fig.savefig(memfile)

    document.add_picture(memfile, width=Inches(6))

document.save('lab02.docx')
