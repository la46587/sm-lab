from docx import Document
from docx.shared import Inches
from io import BytesIO
import cv2
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from scipy.stats import pearsonr, spearmanr


def calcMSE(imgOriginal, imgModified):
    return np.mean((imgOriginal.astype("float") - imgModified.astype("float")) ** 2)


def calcNMSE(imgOriginal, imgModified):
    return calcMSE(imgOriginal, imgModified) / np.var(imgOriginal.astype('float'))


def calcPSNR(imgOriginal, imgModified):
    mse = calcMSE(imgOriginal, imgModified)
    maxPixelValue = 255.0 if imgOriginal.dtype == np.uint8 else 1.0
    psnr = 20 * np.log10(maxPixelValue / np.sqrt(mse))
    return psnr


def calcIF(imgOriginal, imgModified):
    mse = calcMSE(imgOriginal, imgModified)
    errorSum = np.sum(mse)
    signalEnergy = np.sum(imgOriginal.astype('float') ** 2)
    imgFidelity = 1 - (errorSum / signalEnergy)
    return imgFidelity


document = Document()
document.add_heading('Laboratorium 10', 0)

results = pd.read_csv('results.csv', sep=',')
results.index = [f'subject{i + 1}' for i in range(len(results))]
results = results.drop(columns=['Timestamp', 'Imię/pseudonim:'])

resultsKot = results.drop(columns=['2.1/15', '2.2/15', '2.3/15', '2.4/15', '2.5/15', '2.6/15', '2.7/15', '2.8/15',
                                   '2.9/15', '2.10/15', '2.11/15', '2.12/15', '2.13/15', '2.14/15', '2.15/15',
                                   '3.1/15', '3.2/15', '3.3/15', '3.4/15', '3.5/15', '3.6/15', '3.7/15', '3.8/15',
                                   '3.9/15', '3.10/15', '3.11/15', '3.12/15', '3.13/15', '3.14/15', '3.15/15'])
resultsKot = resultsKot.transpose()
refKot = cv2.imread('kot/kot.jpg')
imagesKot = ['kot/1.jpg', 'kot/2.jpg', 'kot/3.jpg', 'kot/4.jpg', 'kot/5.jpg', 'kot/6.jpg', 'kot/7.jpg',
             'kot/8.jpg', 'kot/9.jpg', 'kot/10.jpg', 'kot/11.jpg', 'kot/12.jpg', 'kot/13.jpg', 'kot/14.jpg',
             'kot/15.jpg']

mseKot = []
nmseKot = []
psnrKot = []
ifKot = []

for imagePath in imagesKot:
    image = cv2.imread(imagePath)

    mse = calcMSE(refKot, image)
    mseKot.append(mse)

    nmse = calcNMSE(refKot, image)
    nmseKot.append(nmse)

    psnr = calcPSNR(refKot, image)
    psnrKot.append(psnr)

    IF = calcIF(refKot, image)
    ifKot.append(IF)

resultsKot['MSE'] = mseKot
resultsKot['NMSE'] = nmseKot
resultsKot['PSNR'] = psnrKot
resultsKot['IF'] = ifKot

resultsKotMeans = []

for i in range(1, 16):
    resultsKotMeans.append(round(resultsKot.loc[f'1.{i}/15'].mean(), 2))
resultsKot['MOS Mean'] = resultsKotMeans

correlation_matrix = np.corrcoef(resultsKot[['MOS Mean', 'MSE', 'NMSE', 'PSNR', 'IF']].values.T)
correlation_df = pd.DataFrame(correlation_matrix, index=['MOS Mean', 'MSE', 'NMSE', 'PSNR', 'IF'],
                              columns=['MOS Mean', 'MSE', 'NMSE', 'PSNR', 'IF'])

memfile = BytesIO()
plt.figure(figsize=(10, 8))
sns.heatmap(correlation_df, annot=True, cmap='coolwarm', fmt=".2f")
plt.title('Correlation Matrix Heatmap - Kot')
plt.savefig(memfile, bbox_inches='tight', dpi=500)
document.add_picture(memfile, width=Inches(4.5))
plt.close()

resultsMazda = results.drop(columns=['1.1/15', '1.2/15', '1.3/15', '1.4/15', '1.5/15', '1.6/15', '1.7/15', '1.8/15',
                                     '1.9/15', '1.10/15', '1.11/15', '1.12/15', '1.13/15', '1.14/15', '1.15/15',
                                     '3.1/15', '3.2/15', '3.3/15', '3.4/15', '3.5/15', '3.6/15', '3.7/15', '3.8/15',
                                     '3.9/15', '3.10/15', '3.11/15', '3.12/15', '3.13/15', '3.14/15', '3.15/15'])
resultsMazda = resultsMazda.transpose()
refMazda = cv2.imread('mazda/mazda.jpg')
imagesMazda = ['mazda/1.jpg', 'mazda/2.jpg', 'mazda/3.jpg', 'mazda/4.jpg', 'mazda/5.jpg', 'mazda/6.jpg', 'mazda/7.jpg',
               'mazda/8.jpg', 'mazda/9.jpg', 'mazda/10.jpg', 'mazda/11.jpg', 'mazda/12.jpg', 'mazda/13.jpg', 'mazda/14.jpg',
               'mazda/15.jpg']

mseMazda = []
nmseMazda = []
psnrMazda = []
ifMazda = []

for imagePath in imagesMazda:
    image = cv2.imread(imagePath)

    mse = calcMSE(refMazda, image)
    mseMazda.append(mse)

    nmse = calcNMSE(refMazda, image)
    nmseMazda.append(nmse)

    psnr = calcPSNR(refMazda, image)
    psnrMazda.append(psnr)

    IF = calcIF(refMazda, image)
    ifMazda.append(IF)

resultsMazda['MSE'] = mseMazda
resultsMazda['NMSE'] = nmseMazda
resultsMazda['PSNR'] = psnrMazda
resultsMazda['IF'] = ifMazda

resultsMazdaMeans = []

for i in range(1, 16):
    resultsMazdaMeans.append(round(resultsMazda.loc[f'2.{i}/15'].mean(), 2))
resultsMazda['MOS Mean'] = resultsMazdaMeans

correlation_matrix = np.corrcoef(resultsMazda[['MOS Mean', 'MSE', 'NMSE', 'PSNR', 'IF']].values.T)
correlation_df = pd.DataFrame(correlation_matrix, index=['MOS Mean', 'MSE', 'NMSE', 'PSNR', 'IF'],
                              columns=['MOS Mean', 'MSE', 'NMSE', 'PSNR', 'IF'])

memfile = BytesIO()
plt.figure(figsize=(10, 8))
sns.heatmap(correlation_df, annot=True, cmap='coolwarm', fmt=".2f")
plt.title('Correlation Matrix Heatmap - Mazda')
plt.savefig(memfile, bbox_inches='tight', dpi=500)
document.add_picture(memfile, width=Inches(4.5))
plt.close()

resultsZima = results.drop(columns=['1.1/15', '1.2/15', '1.3/15', '1.4/15', '1.5/15', '1.6/15', '1.7/15', '1.8/15',
                                    '1.9/15', '1.10/15', '1.11/15', '1.12/15', '1.13/15', '1.14/15', '1.15/15',
                                    '2.1/15', '2.2/15', '2.3/15', '2.4/15', '2.5/15', '2.6/15', '2.7/15', '2.8/15',
                                    '2.9/15', '2.10/15', '2.11/15', '2.12/15', '2.13/15', '2.14/15', '2.15/15'])
resultsZima = resultsZima.transpose()
refZima = cv2.imread('zima/zima.png')
refZima = refZima[..., :3]
imagesZima = ['zima/1.jpg', 'zima/2.jpg', 'zima/3.jpg', 'zima/4.jpg', 'zima/5.jpg', 'zima/6.jpg', 'zima/7.jpg',
              'zima/8.jpg', 'zima/9.jpg', 'zima/10.jpg', 'zima/11.jpg', 'zima/12.jpg', 'zima/13.jpg', 'zima/14.jpg',
              'zima/15.jpg']

mseZima = []
nmseZima = []
psnrZima = []
ifZima = []

for imagePath in imagesZima:
    image = cv2.imread(imagePath)

    mse = calcMSE(refZima, image)
    mseZima.append(mse)

    nmse = calcNMSE(refZima, image)
    nmseZima.append(nmse)

    psnr = calcPSNR(refZima, image)
    psnrZima.append(psnr)

    IF = calcIF(refZima, image)
    ifZima.append(IF)

resultsZima['MSE'] = mseZima
resultsZima['NMSE'] = nmseZima
resultsZima['PSNR'] = psnrZima
resultsZima['IF'] = ifZima

resultsZimaMeans = []

for i in range(1, 16):
    resultsZimaMeans.append(round(resultsZima.loc[f'3.{i}/15'].mean(), 2))
resultsZima['MOS Mean'] = resultsZimaMeans

correlation_matrix = np.corrcoef(resultsZima[['MOS Mean', 'MSE', 'NMSE', 'PSNR', 'IF']].values.T)
correlation_df = pd.DataFrame(correlation_matrix, index=['MOS Mean', 'MSE', 'NMSE', 'PSNR', 'IF'],
                              columns=['MOS Mean', 'MSE', 'NMSE', 'PSNR', 'IF'])

memfile = BytesIO()
plt.figure(figsize=(10, 8))
sns.heatmap(correlation_df, annot=True, cmap='coolwarm', fmt=".2f")
plt.title('Correlation Matrix Heatmap - Zima')
plt.savefig(memfile, bbox_inches='tight', dpi=500)
document.add_picture(memfile, width=Inches(4.5))
plt.close()

document.save('lab10.docx')
