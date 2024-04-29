from docx import Document
from docx.shared import Inches
import cv2
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.image as mpimg
from scipy.fftpack import dct, idct
from io import BytesIO


def dct2(a):
    return dct(dct(a.T, norm='ortho').T, norm='ortho')


def idct2(a):
    return idct(idct(a.T, norm='ortho').T, norm='ortho')


class JPEGChromaSubsample:
    def __init__(self, image, ratio):
        self.image = image
        self.ratio = ratio
        self.YCbCr = cv2.cvtColor(image, cv2.COLOR_RGB2YCrCb)
        self.Y, self.Cb, self.Cr = self.YCbCr[:, :, 0], self.YCbCr[:, :, 1], self.YCbCr[:, :, 2]
        self.subsample()

    def subsample(self):
        if self.ratio == "4:2:2":
            self.Cb = self.Cb[:, ::2]
            self.Cr = self.Cr[:, ::2]

    def recover(self):
        YCbCr = np.empty_like(self.YCbCr)
        YCbCr[:, :, 0] = self.Y
        YCbCr[:, :, 1] = np.repeat(self.Cb, 2, axis=1) if self.ratio == "4:2:2" else self.Cb
        YCbCr[:, :, 2] = np.repeat(self.Cr, 2, axis=1) if self.ratio == "4:2:2" else self.Cr
        return cv2.cvtColor(YCbCr, cv2.COLOR_YCrCb2RGB)


def process_channel(channel, Q):
    height, width = channel.shape
    padded_height = (height + 7) // 8 * 8
    padded_width = (width + 7) // 8 * 8
    padded_channel = np.zeros((padded_height, padded_width))
    padded_channel[:height, :width] = channel

    recovered_channel = np.zeros_like(padded_channel)

    for i in range(0, padded_height, 8):
        for j in range(0, padded_width, 8):
            block = padded_channel[i:i + 8, j:j + 8]
            dct_block = dct2(block)
            quantized_block = np.round(dct_block / Q)
            dequantized_block = quantized_block * Q
            recovered_block = idct2(dequantized_block)
            recovered_channel[i:i + 8, j:j + 8] = recovered_block

    return recovered_channel[:height, :width]


class JPEGQuantization:
    def __init__(self, image, QY, QC):
        self.image = image
        if QY is None:
            self.QY = np.ones((8, 8))
        else:
            self.QY = QY
        if QC is None:
            self.QC = np.ones((8, 8))
        else:
            self.QC = QC
        self.YCbCr = cv2.cvtColor(image, cv2.COLOR_RGB2YCrCb)
        self.Y, self.Cb, self.Cr = self.YCbCr[:, :, 0], self.YCbCr[:, :, 1], self.YCbCr[:, :, 2]

    def recover(self):
        Y = process_channel(self.Y, self.QY)
        Cb = process_channel(self.Cb, self.QC)
        Cr = process_channel(self.Cr, self.QC)

        YCbCr = np.stack((Y, Cb, Cr), axis=-1)
        return cv2.cvtColor(YCbCr.astype(np.uint8), cv2.COLOR_YCrCb2RGB)


def calculateCompressionRatio(img, imgEncoded):
    originalSize = img.size
    encodedSize = len(imgEncoded)

    compressionRatio = abs(originalSize) / abs(encodedSize)
    compressionPercentage = 100 * (abs(encodedSize) / abs(originalSize))

    return compressionRatio, compressionPercentage


def encodeRLE(img):
    pixels = img.flatten()
    encoded = []
    count = 1
    prevPixel = pixels[0]

    for pixel in pixels[1:]:
        if np.array_equal(pixel, prevPixel):
            count += 1
        else:
            encoded.append((count, prevPixel))
            prevPixel = pixel
            count = 1
    encoded.append((count, prevPixel))
    return encoded


QY = np.array([
    [16, 11, 10, 16, 24, 40, 51, 61],
    [12, 12, 14, 19, 26, 58, 60, 55],
    [14, 13, 16, 24, 40, 57, 69, 56],
    [14, 17, 22, 29, 51, 87, 80, 62],
    [18, 22, 37, 56, 68, 109, 103, 77],
    [24, 36, 55, 64, 81, 104, 113, 92],
    [49, 64, 78, 87, 103, 121, 120, 101],
    [72, 92, 95, 98, 112, 100, 103, 99]
])

QC = np.array([
    [17, 18, 24, 47, 99, 99, 99, 99],
    [18, 21, 26, 66, 99, 99, 99, 99],
    [24, 26, 56, 99, 99, 99, 99, 99],
    [47, 66, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99],
    [99, 99, 99, 99, 99, 99, 99, 99]
])

document = Document()
document.add_heading('Laboratorium 08', 0)

images = ['frag1_1.jpg', 'frag1_2.jpg', 'frag1_3.jpg', 'frag2_1.jpg', 'frag2_2.jpg', 'frag2_3.jpg', 'frag3_1.jpg',
          'frag3_2.jpg', 'frag3_3.jpg', 'frag4_1.jpg', 'frag4_2.jpg', 'frag4_3.jpg']

for image in images:
    memfile = BytesIO()

    input_image = mpimg.imread(image)
    processed_image = JPEGChromaSubsample(input_image, "4:4:4")
    output_image = processed_image.recover()

    imgEncodedRLE = encodeRLE(output_image.copy())
    compressionRatio, compressionPercentage = calculateCompressionRatio(output_image, imgEncodedRLE)

    fig, axs = plt.subplots(4, 2, sharey='all')
    fig.set_size_inches(9, 13)
    fig.tight_layout()

    axs[0, 0].imshow(input_image)
    axs[0, 0].axis('off')
    axs[0, 1].imshow(output_image)
    axs[0, 1].axis('off')

    input_ycbcr = cv2.cvtColor(input_image, cv2.COLOR_RGB2YCrCb)
    output_ycbcr = cv2.cvtColor(output_image, cv2.COLOR_RGB2YCrCb)

    for i in range(3):
        axs[i + 1, 0].imshow(input_ycbcr[:, :, i], cmap='gray')
        axs[i + 1, 0].axis('off')
        axs[i + 1, 1].imshow(output_ycbcr[:, :, i], cmap='gray')
        axs[i + 1, 1].axis('off')

    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('Redukcja chrominancji - 4:4:4', 1)
    document.add_heading(f'{image} - stopień kompresji: {compressionRatio:.4f} ({compressionPercentage:.2f}%)', 2)
    document.add_picture(memfile, width=Inches(4.3))
    plt.close(fig)

    memfile = BytesIO()

    input_image = mpimg.imread(image)
    processed_image = JPEGChromaSubsample(input_image, "4:2:2")
    output_image = processed_image.recover()

    imgEncodedRLE = encodeRLE(output_image.copy())
    compressionRatio, compressionPercentage = calculateCompressionRatio(output_image, imgEncodedRLE)

    fig, axs = plt.subplots(4, 2, sharey='all')
    fig.set_size_inches(9, 13)
    fig.tight_layout()

    axs[0, 0].imshow(input_image)
    axs[0, 0].axis('off')
    axs[0, 1].imshow(output_image)
    axs[0, 1].axis('off')

    input_ycbcr = cv2.cvtColor(input_image, cv2.COLOR_RGB2YCrCb)
    output_ycbcr = cv2.cvtColor(output_image, cv2.COLOR_RGB2YCrCb)

    for i in range(3):
        axs[i + 1, 0].imshow(input_ycbcr[:, :, i], cmap='gray')
        axs[i + 1, 0].axis('off')
        axs[i + 1, 1].imshow(output_ycbcr[:, :, i], cmap='gray')
        axs[i + 1, 1].axis('off')

    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('Redukcja chrominancji - 4:2:2', 1)
    document.add_heading(f'{image} - stopień kompresji: {compressionPercentage:.4f} ({compressionPercentage:.2f}%)', 2)
    document.add_picture(memfile, width=Inches(4.3))
    plt.close(fig)

    memfile = BytesIO()

    input_image = mpimg.imread(image)
    processed_image = JPEGQuantization(input_image, QY, QC)
    output_image = processed_image.recover()

    imgEncodedRLE = encodeRLE(output_image.copy())
    compressionRatio, compressionPercentage = calculateCompressionRatio(output_image, imgEncodedRLE)

    fig, axs = plt.subplots(4, 2, sharey='all')
    fig.set_size_inches(9, 13)
    fig.tight_layout()

    axs[0, 0].imshow(input_image)
    axs[0, 0].axis('off')
    axs[0, 1].imshow(output_image)
    axs[0, 1].axis('off')

    input_ycbcr = cv2.cvtColor(input_image, cv2.COLOR_RGB2YCrCb)
    output_ycbcr = cv2.cvtColor(output_image, cv2.COLOR_RGB2YCrCb)

    for i in range(3):
        axs[i + 1, 0].imshow(input_ycbcr[:, :, i], cmap='gray')
        axs[i + 1, 0].axis('off')
        axs[i + 1, 1].imshow(output_ycbcr[:, :, i], cmap='gray')
        axs[i + 1, 1].axis('off')

    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('Kwantyzacja - tablice kwantyzujące', 1)
    document.add_heading(f'{image} - stopień kompresji: {compressionPercentage:.4f} ({compressionPercentage:.2f}%)', 2)
    document.add_picture(memfile, width=Inches(4.3))
    plt.close(fig)

    memfile = BytesIO()

    input_image = mpimg.imread(image)
    processed_image = JPEGQuantization(input_image, None, None)
    output_image = processed_image.recover()

    imgEncodedRLE = encodeRLE(output_image.copy())
    compressionRatio, compressionPercentage = calculateCompressionRatio(output_image, imgEncodedRLE)

    fig, axs = plt.subplots(4, 2, sharey='all')
    fig.set_size_inches(9, 13)
    fig.tight_layout()

    axs[0, 0].imshow(input_image)
    axs[0, 0].axis('off')
    axs[0, 1].imshow(output_image)
    axs[0, 1].axis('off')

    input_ycbcr = cv2.cvtColor(input_image, cv2.COLOR_RGB2YCrCb)
    output_ycbcr = cv2.cvtColor(output_image, cv2.COLOR_RGB2YCrCb)

    for i in range(3):
        axs[i + 1, 0].imshow(input_ycbcr[:, :, i], cmap='gray')
        axs[i + 1, 0].axis('off')
        axs[i + 1, 1].imshow(output_ycbcr[:, :, i], cmap='gray')
        axs[i + 1, 1].axis('off')

    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('Kwantyzacja - tablica jedynek', 1)
    document.add_heading(f'{image} - stopień kompresji: {compressionPercentage:.4f} ({compressionPercentage:.2f}%)', 2)
    document.add_picture(memfile, width=Inches(4.3))
    plt.close(fig)

document.save('lab08.docx')
