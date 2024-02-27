from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
import scipy.fftpack
import sounddevice as sd
import soundfile as sf
from io import BytesIO


def plotAudio(Axs, Signal, Fs, TimeMargin=[0, 0.02]):
    plt.subplot(2, 1, 1)
    plt.xlim(TimeMargin)
    Axs[0] = plt.plot(np.arange(0, Signal.shape[0]) / Fs, Signal)
    plt.xlabel('czas trwania [s]')
    plt.subplot(2, 1, 2)
    yf = scipy.fftpack.fft(Signal, fsize)
    Axs[1] = plt.plot(np.arange(0, Fs / 2, Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])))
    plt.xlabel('częstotliwość [Hz]')
    plt.ylabel('[dB]')

    amplitude = max(Signal / 1e9) - min(Signal / 1e9)
    maxdBsample = np.argmax(20 * np.log10(np.abs(yf[:fsize // 2])))
    maxdB = np.arange(0, Fs / 2, Fs / fsize)[maxdBsample]

    return amplitude, maxdB


np.seterr(divide='ignore')
document = Document()
document.add_heading('Laboratorium 01 - Zadanie 3', 0)

files = ['sin_60Hz.wav', 'sin_440Hz.wav', 'sin_8000Hz.wav']
fsizes = [2**8, 2**12, 2**16]
for file in files:
    document.add_heading('Plik - {}'.format(file), 2)
    for fsize in fsizes:
        document.add_heading('Fsize = {}'.format(fsize), 3)
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))

        data, fs = sf.read(file, dtype=np.int32)
        amplitude, maxdB = plotAudio(axs, data, fs)

        fig.suptitle('Time margin {}'.format([0, 0.02]))
        fig.tight_layout(pad=1.5)
        memfile = BytesIO()
        fig.savefig(memfile)

        document.add_picture(memfile, width=Inches(6))

        memfile.close()

        document.add_paragraph('amplituda = {:.2f}e9'.format(amplitude))
        document.add_paragraph('częstotliwość o najwyższej wartości = {:.2f} Hz'.format(maxdB))

document.save('lab01.docx')
