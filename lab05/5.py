from docx import Document
from docx.shared import Inches
from io import BytesIO
from scipy.interpolate import interp1d
import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import sounddevice as sd
import soundfile as sf


def quantize(signal, targetBits):
    targetMax = 2 ** targetBits - 1
    quantizedSignal = np.round(signal * (targetMax / 2)) / (targetMax / 2)
    return quantizedSignal


def decimate(signal, n):
    decimatedSignal = signal[::n]
    return decimatedSignal


def interpolate(signal, rate, method):
    oldSignal = np.arange(0, len(signal))
    newSignal = np.linspace(0, len(signal) - 1, rate)

    if method == 'linear':
        interpolatedSignal = np.interp(newSignal, oldSignal, signal)
    elif method == 'cubic':
        func = interp1d(oldSignal, signal, kind='cubic')
        interpolatedSignal = func(newSignal)
    else:
        return print('invalid interpolation method')
    return interpolatedSignal


np.seterr(divide='ignore')

document = Document()
document.add_heading('Laboratorium 05', 0)

sin60Hz, sin60HzFs = sf.read('sin_60Hz.wav', dtype='float32')
sin440Hz, sin440HzFs = sf.read('sin_440Hz.wav', dtype='float32')
sin8000Hz, sin8000HzFs = sf.read('sin_8000Hz.wav', dtype='float32')
sinCombined, sinCombinedFs = sf.read('sin_Combined.wav', dtype='float32')

fsize = 2**16
signals = [sin60Hz, sin440Hz, sin8000Hz, sinCombined]
signalsFs = [sin60HzFs, sin440HzFs, sin8000HzFs, sinCombinedFs]
signalLabels = ['60Hz', '440Hz', '8000Hz', 'combined']

# kwantyzacja
quantizeBits = [4, 8, 16, 24]
for signal, signalFs, label in zip(signals, signalsFs, signalLabels):
    originalSignal = signal.copy()

    signal4bits = quantize(originalSignal, 4)
    signal8bits = quantize(originalSignal, 8)
    signal16bits = quantize(originalSignal, 16)
    signal24bits = quantize(originalSignal, 24)

    if label == '60Hz':
        timeMargin = [0, 0.1]
    elif label == '440Hz':
        timeMargin = [0, 0.01]
    elif label == '8000Hz':
        timeMargin = [0, 0.001]
    else:
        timeMargin = [0, 0.002]

    fig, axs = plt.subplots(2, 4, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("quantization: {}".format(label))

    axs[0, 0].plot(np.arange(0, signal4bits.shape[0]) / signalFs, signal4bits, color='black')
    axs[0, 0].set_title('4 bits')
    axs[0, 0].set_xlim(timeMargin)
    axs[0, 0].set_xticks(timeMargin)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, signal8bits.shape[0]) / signalFs, signal8bits, color='black')
    axs[0, 1].set_title('8 bits')
    axs[0, 1].set_xlim(timeMargin)
    axs[0, 1].set_xticks(timeMargin)
    axs[0, 1].set_xlabel('czas trwania [s]')

    axs[0, 2].plot(np.arange(0, signal16bits.shape[0]) / signalFs, signal16bits, color='black')
    axs[0, 2].set_title('16 bits')
    axs[0, 2].set_xlim(timeMargin)
    axs[0, 2].set_xticks(timeMargin)
    axs[0, 2].set_xlabel('czas trwania [s]')

    axs[0, 3].plot(np.arange(0, signal24bits.shape[0]) / signalFs, signal24bits, color='black')
    axs[0, 3].set_title('24 bits')
    axs[0, 3].set_xlim(timeMargin)
    axs[0, 3].set_xticks(timeMargin)
    axs[0, 3].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(signal4bits, fsize)
    axs[1, 0].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal8bits, fsize)
    axs[1, 1].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal16bits, fsize)
    axs[1, 2].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 2].set_xlabel('częstotliwość [Hz]')
    axs[1, 2].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal24bits, fsize)
    axs[1, 3].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 3].set_xlabel('częstotliwość [Hz]')
    axs[1, 3].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('kwantyzacja - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

# decymacja
decimatedSteps = [2, 4, 6, 10, 24]
for signal, signalFs, label in zip(signals, signalsFs, signalLabels):
    originalSignal = signal.copy()

    signal2steps = decimate(originalSignal, 2)
    signal4steps = decimate(originalSignal, 4)
    signal6steps = decimate(originalSignal, 6)
    signal10steps = decimate(originalSignal, 10)
    signal24steps = decimate(originalSignal, 24)

    if label == '60Hz':
        timeMargin = [0, 0.04]
    elif label == '440Hz':
        timeMargin = [0, 0.005]
    elif label == '8000Hz':
        timeMargin = [0, 0.002]
    else:
        timeMargin = [0, 0.002]

    fig, axs = plt.subplots(4, 3, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("decimation: {}".format(label))

    axs[0, 0].plot(np.arange(0, signal2steps.shape[0]) / signalFs, signal2steps, color='black')
    axs[0, 0].set_title('2 steps')
    axs[0, 0].set_xlim(timeMargin)
    axs[0, 0].set_xticks(timeMargin)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, signal4steps.shape[0]) / signalFs, signal4steps, color='black')
    axs[0, 1].set_title('4 steps')
    axs[0, 1].set_xlim(timeMargin)
    axs[0, 1].set_xticks(timeMargin)
    axs[0, 1].set_xlabel('czas trwania [s]')

    axs[0, 2].plot(np.arange(0, signal6steps.shape[0]) / signalFs, signal6steps, color='black')
    axs[0, 2].set_title('6 steps')
    axs[0, 2].set_xlim(timeMargin)
    axs[0, 2].set_xticks(timeMargin)
    axs[0, 2].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(signal2steps, fsize)
    axs[1, 0].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal4steps, fsize)
    axs[1, 1].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal6steps, fsize)
    axs[1, 2].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 2].set_xlabel('częstotliwość [Hz]')
    axs[1, 2].set_ylabel('[dB]')

    axs[2, 0].plot(np.arange(0, signal10steps.shape[0]) / signalFs, signal10steps, color='black')
    axs[2, 0].set_title('10 steps')
    axs[2, 0].set_xlim(timeMargin)
    axs[2, 0].set_xticks(timeMargin)
    axs[2, 0].set_xlabel('czas trwania [s]')

    axs[2, 1].plot(np.arange(0, signal24steps.shape[0]) / signalFs, signal24steps, color='black')
    axs[2, 1].set_title('24 steps')
    axs[2, 1].set_xlim(timeMargin)
    axs[2, 1].set_xticks(timeMargin)
    axs[2, 1].set_xlabel('czas trwania [s]')

    axs[2, 2].set_visible(False)

    yf = scipy.fftpack.fft(signal10steps, fsize)
    axs[3, 0].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 0].set_xlabel('częstotliwość [Hz]')
    axs[3, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal24steps, fsize)
    axs[3, 1].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 1].set_xlabel('częstotliwość [Hz]')
    axs[3, 1].set_ylabel('[dB]')

    axs[3, 2].set_visible(False)

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('decymacja - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

# interpolacja liniowa
rates = [2000, 4000, 8000, 11999, 16000, 16953, 24000, 41000]
for signal, signalFs, label in zip(signals, signalsFs, signalLabels):
    originalSignal = signal.copy()

    signal2000 = interpolate(signal, 2000, 'linear')
    signal4000 = interpolate(signal, 4000, 'linear')
    signal8000 = interpolate(signal, 8000, 'linear')
    signal11999 = interpolate(signal, 11999, 'linear')
    signal16000 = interpolate(signal, 16000, 'linear')
    signal16953 = interpolate(signal, 16953, 'linear')
    signal24000 = interpolate(signal, 24000, 'linear')
    signal41000 = interpolate(signal, 41000, 'linear')

    fig, axs = plt.subplots(4, 4, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("linear interpolation: {}".format(label))

    axs[0, 0].plot(np.arange(0, signal2000.shape[0]) / signalFs, signal2000, color='black')
    axs[0, 0].set_title('2000')
    axs[0, 0].set_xlim(0, 0.001)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, signal4000.shape[0]) / signalFs, signal4000, color='black')
    axs[0, 1].set_title('4000')
    axs[0, 1].set_xlim(0, 0.001)
    axs[0, 1].set_xlabel('czas trwania [s]')

    axs[0, 2].plot(np.arange(0, signal8000.shape[0]) / signalFs, signal8000, color='black')
    axs[0, 2].set_title('8000')
    axs[0, 2].set_xlim(0, 0.001)
    axs[0, 2].set_xlabel('czas trwania [s]')

    axs[0, 3].plot(np.arange(0, signal11999.shape[0]) / signalFs, signal11999, color='black')
    axs[0, 3].set_title('11999')
    axs[0, 3].set_xlim(0, 0.001)
    axs[0, 3].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(signal2000, fsize)
    axs[1, 0].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal4000, fsize)
    axs[1, 1].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal8000, fsize)
    axs[1, 2].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 2].set_xlabel('częstotliwość [Hz]')
    axs[1, 2].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal11999, fsize)
    axs[1, 3].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 3].set_xlabel('częstotliwość [Hz]')
    axs[1, 3].set_ylabel('[dB]')

    axs[2, 0].plot(np.arange(0, signal16000.shape[0]) / signalFs, signal16000, color='black')
    axs[2, 0].set_title('16000')
    axs[2, 0].set_xlim(0, 0.001)
    axs[2, 0].set_xlabel('czas trwania [s]')

    axs[2, 1].plot(np.arange(0, signal16953.shape[0]) / signalFs, signal16953, color='black')
    axs[2, 1].set_title('16953')
    axs[2, 1].set_xlim(0, 0.001)
    axs[2, 1].set_xlabel('czas trwania [s]')

    axs[2, 2].plot(np.arange(0, signal24000.shape[0]) / signalFs, signal24000, color='black')
    axs[2, 2].set_title('24000')
    axs[2, 2].set_xlim(0, 0.001)
    axs[2, 2].set_xlabel('czas trwania [s]')

    axs[2, 3].plot(np.arange(0, signal41000.shape[0]) / signalFs, signal41000, color='black')
    axs[2, 3].set_title('41000')
    axs[2, 3].set_xlim(0, 0.001)
    axs[2, 3].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(signal16000, fsize)
    axs[3, 0].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 0].set_xlabel('częstotliwość [Hz]')
    axs[3, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal16953, fsize)
    axs[3, 1].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 1].set_xlabel('częstotliwość [Hz]')
    axs[3, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal24000, fsize)
    axs[3, 2].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 2].set_xlabel('częstotliwość [Hz]')
    axs[3, 2].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal41000, fsize)
    axs[3, 3].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 3].set_xlabel('częstotliwość [Hz]')
    axs[3, 3].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('interpolacja liniowa - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

# interpolacja nieliniowa
for signal, signalFs, label in zip(signals, signalsFs, signalLabels):
    originalSignal = signal.copy()

    signal2000 = interpolate(signal, 2000, 'cubic')
    signal4000 = interpolate(signal, 4000, 'cubic')
    signal8000 = interpolate(signal, 8000, 'cubic')
    signal11999 = interpolate(signal, 11999, 'cubic')
    signal16000 = interpolate(signal, 16000, 'cubic')
    signal16953 = interpolate(signal, 16953, 'cubic')
    signal24000 = interpolate(signal, 24000, 'cubic')
    signal41000 = interpolate(signal, 41000, 'cubic')

    fig, axs = plt.subplots(4, 4, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("nonlinear interpolation: {}".format(label))

    axs[0, 0].plot(np.arange(0, signal2000.shape[0]) / signalFs, signal2000, color='black')
    axs[0, 0].set_title('2000')
    axs[0, 0].set_xlim(0, 0.001)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, signal4000.shape[0]) / signalFs, signal4000, color='black')
    axs[0, 1].set_title('4000')
    axs[0, 1].set_xlim(0, 0.001)
    axs[0, 1].set_xlabel('czas trwania [s]')

    axs[0, 2].plot(np.arange(0, signal8000.shape[0]) / signalFs, signal8000, color='black')
    axs[0, 2].set_title('8000')
    axs[0, 2].set_xlim(0, 0.001)
    axs[0, 2].set_xlabel('czas trwania [s]')

    axs[0, 3].plot(np.arange(0, signal11999.shape[0]) / signalFs, signal11999, color='black')
    axs[0, 3].set_title('11999')
    axs[0, 3].set_xlim(0, 0.001)
    axs[0, 3].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(signal2000, fsize)
    axs[1, 0].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal4000, fsize)
    axs[1, 1].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal8000, fsize)
    axs[1, 2].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 2].set_xlabel('częstotliwość [Hz]')
    axs[1, 2].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal11999, fsize)
    axs[1, 3].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 3].set_xlabel('częstotliwość [Hz]')
    axs[1, 3].set_ylabel('[dB]')

    axs[2, 0].plot(np.arange(0, signal16000.shape[0]) / signalFs, signal16000, color='black')
    axs[2, 0].set_title('16000')
    axs[2, 0].set_xlim(0, 0.001)
    axs[2, 0].set_xlabel('czas trwania [s]')

    axs[2, 1].plot(np.arange(0, signal16953.shape[0]) / signalFs, signal16953, color='black')
    axs[2, 1].set_title('16953')
    axs[2, 1].set_xlim(0, 0.001)
    axs[2, 1].set_xlabel('czas trwania [s]')

    axs[2, 2].plot(np.arange(0, signal24000.shape[0]) / signalFs, signal24000, color='black')
    axs[2, 2].set_title('24000')
    axs[2, 2].set_xlim(0, 0.001)
    axs[2, 2].set_xlabel('czas trwania [s]')

    axs[2, 3].plot(np.arange(0, signal41000.shape[0]) / signalFs, signal41000, color='black')
    axs[2, 3].set_title('41000')
    axs[2, 3].set_xlim(0, 0.001)
    axs[2, 3].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(signal16000, fsize)
    axs[3, 0].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 0].set_xlabel('częstotliwość [Hz]')
    axs[3, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal16953, fsize)
    axs[3, 1].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 1].set_xlabel('częstotliwość [Hz]')
    axs[3, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal24000, fsize)
    axs[3, 2].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 2].set_xlabel('częstotliwość [Hz]')
    axs[3, 2].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal41000, fsize)
    axs[3, 3].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 3].set_xlabel('częstotliwość [Hz]')
    axs[3, 3].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('interpolacja nieliniowa - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

singHigh, singHighFs = sf.read('sing_high2.wav', dtype='float32')
singMedium, singMediumFs = sf.read('sing_medium1.wav', dtype='float32')
singLow, singLowFs = sf.read('sing_low1.wav', dtype='float32')

sings = [singHigh, singMedium, singLow]
singsFs = [singHighFs, singMediumFs, singLowFs]
singLabels = ['high', 'medium', 'low']

# kwantyzacja
for sing, singFs, label in zip(sings, singsFs, singLabels):
    originalSing = sing.copy()

    sing4bits = quantize(originalSing, 4)
    sing8bits = quantize(originalSing, 8)

    if label == 'high':
        timeMargin = [2, 2.01]
    elif label == 'medium':
        timeMargin = [2, 2.01]
    else:
        timeMargin = [2, 3]

    fig, axs = plt.subplots(2, 2, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("quantization: {}".format(label))

    axs[0, 0].plot(np.arange(0, sing4bits.shape[0]) / singFs, sing4bits, color='black')
    axs[0, 0].set_title('4 bits')
    axs[0, 0].set_xlim(timeMargin)
    axs[0, 0].set_xticks(timeMargin)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, sing8bits.shape[0]) / singFs, sing8bits, color='black')
    axs[0, 1].set_title('8 bits')
    axs[0, 1].set_xlim(timeMargin)
    axs[0, 1].set_xticks(timeMargin)
    axs[0, 1].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(sing4bits, fsize)
    axs[1, 0].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing8bits, fsize)
    axs[1, 1].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('kwantyzacja - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

# decymacja
for sing, singFs, label in zip(sings, singsFs, singLabels):
    originalSing = sing.copy()

    sing4steps = decimate(originalSing, 4)
    sing6steps = decimate(originalSing, 6)
    sing10steps = decimate(originalSing, 10)
    sing24steps = decimate(originalSing, 24)

    timeMargin = [0.01, 0.015]

    fig, axs = plt.subplots(4, 2, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("quantization: {}".format(label))

    axs[0, 0].plot(np.arange(0, sing4steps.shape[0]) / singFs, sing4steps, color='black')
    axs[0, 0].set_title('4 steps')
    axs[0, 0].set_xlim(timeMargin)
    axs[0, 0].set_xticks(timeMargin)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, sing6steps.shape[0]) / singFs, sing6steps, color='black')
    axs[0, 1].set_title('6 steps')
    axs[0, 1].set_xlim(timeMargin)
    axs[0, 1].set_xticks(timeMargin)
    axs[0, 1].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(sing4steps, fsize)
    axs[1, 0].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing6steps, fsize)
    axs[1, 1].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    axs[2, 0].plot(np.arange(0, sing10steps.shape[0]) / singFs, sing10steps, color='black')
    axs[2, 0].set_title('10 steps')
    axs[2, 0].set_xlim(timeMargin)
    axs[2, 0].set_xticks(timeMargin)
    axs[2, 0].set_xlabel('czas trwania [s]')

    axs[2, 1].plot(np.arange(0, sing24steps.shape[0]) / singFs, sing24steps, color='black')
    axs[2, 1].set_title('24 steps')
    axs[2, 1].set_xlim(timeMargin)
    axs[2, 1].set_xticks(timeMargin)
    axs[2, 1].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(sing10steps, fsize)
    axs[3, 0].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 0].set_xlabel('częstotliwość [Hz]')
    axs[3, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing24steps, fsize)
    axs[3, 1].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 1].set_xlabel('częstotliwość [Hz]')
    axs[3, 1].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('decymacja - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

# interpolacja liniowa
for sing, singFs, label in zip(sings, singsFs, singLabels):
    originalSing = sing.copy()

    sing4000 = interpolate(originalSing, 4000, 'linear')
    sing8000 = interpolate(originalSing, 8000, 'linear')
    sing11999 = interpolate(originalSing, 11999, 'linear')
    sing16000 = interpolate(originalSing, 16000, 'linear')
    sing16953 = interpolate(originalSing, 16953, 'linear')

    timeMargin = [0.001, 0.003]

    fig, axs = plt.subplots(4, 3, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("interpolation: {}".format(label))

    axs[0, 0].plot(np.arange(0, sing4000.shape[0]) / singFs, sing4000, color='black')
    axs[0, 0].set_title('4 steps')
    axs[0, 0].set_xlim(timeMargin)
    axs[0, 0].set_xticks(timeMargin)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, sing8000.shape[0]) / singFs, sing8000, color='black')
    axs[0, 1].set_title('6 steps')
    axs[0, 1].set_xlim(timeMargin)
    axs[0, 1].set_xticks(timeMargin)
    axs[0, 1].set_xlabel('czas trwania [s]')

    axs[0, 2].plot(np.arange(0, sing11999.shape[0]) / singFs, sing11999, color='black')
    axs[0, 2].set_title('6 steps')
    axs[0, 2].set_xlim(timeMargin)
    axs[0, 2].set_xticks(timeMargin)
    axs[0, 2].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(sing4000, fsize)
    axs[1, 0].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing8000, fsize)
    axs[1, 1].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing11999, fsize)
    axs[1, 2].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 2].set_xlabel('częstotliwość [Hz]')
    axs[1, 2].set_ylabel('[dB]')

    axs[2, 0].plot(np.arange(0, sing16000.shape[0]) / singFs, sing16000, color='black')
    axs[2, 0].set_title('10 steps')
    axs[2, 0].set_xlim(timeMargin)
    axs[2, 0].set_xticks(timeMargin)
    axs[2, 0].set_xlabel('czas trwania [s]')

    axs[2, 1].plot(np.arange(0, sing16953.shape[0]) / singFs, sing16953, color='black')
    axs[2, 1].set_title('24 steps')
    axs[2, 1].set_xlim(timeMargin)
    axs[2, 1].set_xticks(timeMargin)
    axs[2, 1].set_xlabel('czas trwania [s]')

    axs[2, 2].set_visible(False)

    yf = scipy.fftpack.fft(sing16000, fsize)
    axs[3, 0].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 0].set_xlabel('częstotliwość [Hz]')
    axs[3, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing16953, fsize)
    axs[3, 1].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 1].set_xlabel('częstotliwość [Hz]')
    axs[3, 1].set_ylabel('[dB]')

    axs[3, 2].set_visible(False)

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('interpolacja liniowa - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

# interpolacja nieliniowa
for sing, singFs, label in zip(sings, singsFs, singLabels):
    originalSing = sing.copy()

    sing4000 = interpolate(originalSing, 4000, 'cubic')
    sing8000 = interpolate(originalSing, 8000, 'cubic')
    sing11999 = interpolate(originalSing, 11999, 'cubic')
    sing16000 = interpolate(originalSing, 16000, 'cubic')
    sing16953 = interpolate(originalSing, 16953, 'cubic')

    timeMargin = [0.001, 0.003]

    fig, axs = plt.subplots(4, 3, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("interpolation: {}".format(label))

    axs[0, 0].plot(np.arange(0, sing4000.shape[0]) / singFs, sing4000, color='black')
    axs[0, 0].set_title('4 steps')
    axs[0, 0].set_xlim(timeMargin)
    axs[0, 0].set_xticks(timeMargin)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, sing8000.shape[0]) / singFs, sing8000, color='black')
    axs[0, 1].set_title('6 steps')
    axs[0, 1].set_xlim(timeMargin)
    axs[0, 1].set_xticks(timeMargin)
    axs[0, 1].set_xlabel('czas trwania [s]')

    axs[0, 2].plot(np.arange(0, sing11999.shape[0]) / singFs, sing11999, color='black')
    axs[0, 2].set_title('6 steps')
    axs[0, 2].set_xlim(timeMargin)
    axs[0, 2].set_xticks(timeMargin)
    axs[0, 2].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(sing4000, fsize)
    axs[1, 0].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing8000, fsize)
    axs[1, 1].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing11999, fsize)
    axs[1, 2].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 2].set_xlabel('częstotliwość [Hz]')
    axs[1, 2].set_ylabel('[dB]')

    axs[2, 0].plot(np.arange(0, sing16000.shape[0]) / singFs, sing16000, color='black')
    axs[2, 0].set_title('10 steps')
    axs[2, 0].set_xlim(timeMargin)
    axs[2, 0].set_xticks(timeMargin)
    axs[2, 0].set_xlabel('czas trwania [s]')

    axs[2, 1].plot(np.arange(0, sing16953.shape[0]) / singFs, sing16953, color='black')
    axs[2, 1].set_title('24 steps')
    axs[2, 1].set_xlim(timeMargin)
    axs[2, 1].set_xticks(timeMargin)
    axs[2, 1].set_xlabel('czas trwania [s]')

    axs[2, 2].set_visible(False)

    yf = scipy.fftpack.fft(sing16000, fsize)
    axs[3, 0].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 0].set_xlabel('częstotliwość [Hz]')
    axs[3, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing16953, fsize)
    axs[3, 1].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[3, 1].set_xlabel('częstotliwość [Hz]')
    axs[3, 1].set_ylabel('[dB]')

    axs[3, 2].set_visible(False)

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('interpolacja nieliniowa - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

document.save('lab05.docx')
