from docx import Document
from docx.shared import Inches
from io import BytesIO
from scipy.interpolate import interp1d
import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import sounddevice as sd
import soundfile as sf


def quantize(signal, targetBits):
    targetMax = 2 ** targetBits - 1
    quantizedSignal = np.round(signal * (targetMax / 2)) / (targetMax / 2)
    return quantizedSignal


def decimate(signal, fs, n):
    decimatedSignal = signal[::n]
    fs = fs / n
    return decimatedSignal, fs


def interpolate(signal, fs, rate, method):
    oldSignal = np.arange(0, len(signal)) / fs
    fs = fs * rate
    newSignal = np.linspace(0, oldSignal[-1], int(len(signal) * rate))

    if method == 'linear':
        interpolatedSignal = np.interp(newSignal, oldSignal, signal)
    elif method == 'cubic':
        func = interp1d(oldSignal, signal, kind='cubic')
        interpolatedSignal = func(newSignal)
    else:
        return print('invalid interpolation method')
    return interpolatedSignal, fs


np.seterr(divide='ignore')

document = Document()
document.add_heading('Laboratorium 05', 0)

sin60Hz, sin60HzFs = sf.read('sin_60Hz.wav', dtype='float32')
sin440Hz, sin440HzFs = sf.read('sin_440Hz.wav', dtype='float32')
sin8000Hz, sin8000HzFs = sf.read('sin_8000Hz.wav', dtype='float32')
sinCombined, sinCombinedFs = sf.read('sin_Combined.wav', dtype='float32')

fsize = 2**16
signals = [sin60Hz, sin440Hz, sin8000Hz, sinCombined]
signalsFs = [sin60HzFs, sin440HzFs, sin8000HzFs, sinCombinedFs]
signalLabels = ['60Hz', '440Hz', '8000Hz', 'combined']

# kwantyzacja
quantizeBits = [4, 8, 16, 24]
for signal, signalFs, label in zip(signals, signalsFs, signalLabels):
    originalSignal = signal.copy()

    signal4bits = quantize(originalSignal, 4)
    signal8bits = quantize(originalSignal, 8)
    signal16bits = quantize(originalSignal, 16)
    signal24bits = quantize(originalSignal, 24)

    if label == '60Hz':
        timeMargin = [0, 0.1]
    elif label == '440Hz':
        timeMargin = [0, 0.01]
    elif label == '8000Hz':
        timeMargin = [0, 0.001]
    else:
        timeMargin = [0, 0.002]

    fig, axs = plt.subplots(2, 4, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("quantization: {}".format(label))

    axs[0, 0].plot(np.arange(0, signal4bits.shape[0]) / signalFs, signal4bits, color='black')
    axs[0, 0].set_title('4 bits')
    axs[0, 0].set_xlim(timeMargin)
    axs[0, 0].set_xticks(timeMargin)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, signal8bits.shape[0]) / signalFs, signal8bits, color='black')
    axs[0, 1].set_title('8 bits')
    axs[0, 1].set_xlim(timeMargin)
    axs[0, 1].set_xticks(timeMargin)
    axs[0, 1].set_xlabel('czas trwania [s]')

    axs[0, 2].plot(np.arange(0, signal16bits.shape[0]) / signalFs, signal16bits, color='black')
    axs[0, 2].set_title('16 bits')
    axs[0, 2].set_xlim(timeMargin)
    axs[0, 2].set_xticks(timeMargin)
    axs[0, 2].set_xlabel('czas trwania [s]')

    axs[0, 3].plot(np.arange(0, signal24bits.shape[0]) / signalFs, signal24bits, color='black')
    axs[0, 3].set_title('24 bits')
    axs[0, 3].set_xlim(timeMargin)
    axs[0, 3].set_xticks(timeMargin)
    axs[0, 3].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(signal4bits, fsize)
    axs[1, 0].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal8bits, fsize)
    axs[1, 1].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal16bits, fsize)
    axs[1, 2].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 2].set_xlabel('częstotliwość [Hz]')
    axs[1, 2].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal24bits, fsize)
    axs[1, 3].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 3].set_xlabel('częstotliwość [Hz]')
    axs[1, 3].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('kwantyzacja - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

# decymacja
decimatedSteps = [2, 4, 6, 10, 24]
for signal, signalFs, label in zip(signals, signalsFs, signalLabels):
    originalSignal = signal.copy()

    signal2steps, signal2stepsFs = decimate(originalSignal, signalFs, 2)
    signal4steps, signal4stepsFs = decimate(originalSignal, signalFs,  4)
    signal6steps, signal6stepsFs = decimate(originalSignal, signalFs, 6)
    signal10steps, signal10stepsFs = decimate(originalSignal, signalFs, 10)
    signal24steps, signal24stepsFs = decimate(originalSignal, signalFs, 24)

    if label == '60Hz':
        timeMargin = [0, 0.1]
    elif label == '440Hz':
        timeMargin = [0, 0.02]
    elif label == '8000Hz':
        timeMargin = [0, 0.005]
    else:
        timeMargin = [0, 0.01]

    fig, axs = plt.subplots(4, 3, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("decimation: {}".format(label))

    axs[0, 0].plot(np.arange(0, originalSignal.shape[0]) / signalFs, originalSignal, color='black')
    axs[0, 0].set_title('original')
    axs[0, 0].set_xlim(timeMargin)
    axs[0, 0].set_xticks(timeMargin)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, signal2steps.shape[0]) / signal2stepsFs, signal2steps, color='black')
    axs[0, 1].set_title('2 steps')
    axs[0, 1].set_xlim(timeMargin)
    axs[0, 1].set_xticks(timeMargin)
    axs[0, 1].set_xlabel('czas trwania [s]')

    axs[0, 2].plot(np.arange(0, signal4steps.shape[0]) / signal4stepsFs, signal4steps, color='black')
    axs[0, 2].set_title('4 steps')
    axs[0, 2].set_xlim(timeMargin)
    axs[0, 2].set_xticks(timeMargin)
    axs[0, 2].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(originalSignal, fsize)
    axs[1, 0].plot(np.arange(0, signalFs / 2, signalFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal2steps, fsize)
    axs[1, 1].plot(np.arange(0, signal2stepsFs / 2, signal2stepsFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal4steps, fsize)
    axs[1, 2].plot(np.arange(0, signal4stepsFs / 2, signal4stepsFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 2].set_xlabel('częstotliwość [Hz]')
    axs[1, 2].set_ylabel('[dB]')

    axs[2, 0].plot(np.arange(0, signal6steps.shape[0]) / signal6stepsFs, signal6steps, color='black')
    axs[2, 0].set_title('6 steps')
    axs[2, 0].set_xlim(timeMargin)
    axs[2, 0].set_xticks(timeMargin)
    axs[2, 0].set_xlabel('czas trwania [s]')

    axs[2, 1].plot(np.arange(0, signal10steps.shape[0]) / signal10stepsFs, signal10steps, color='black')
    axs[2, 1].set_title('10 steps')
    axs[2, 1].set_xlim(timeMargin)
    axs[2, 1].set_xticks(timeMargin)
    axs[2, 1].set_xlabel('czas trwania [s]')

    axs[2, 2].plot(np.arange(0, signal24steps.shape[0]) / signal24stepsFs, signal24steps, color='black')
    axs[2, 2].set_title('24 steps')
    axs[2, 2].set_xlim(timeMargin)
    axs[2, 2].set_xticks(timeMargin)
    axs[2, 2].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(signal6steps, fsize)
    axs[3, 0].plot(np.arange(0, signal6stepsFs / 2, signal6stepsFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 0].set_xlabel('częstotliwość [Hz]')
    axs[3, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal10steps, fsize)
    axs[3, 1].plot(np.arange(0, signal10stepsFs / 2, signal10stepsFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 1].set_xlabel('częstotliwość [Hz]')
    axs[3, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal24steps, fsize)
    axs[3, 2].plot(np.arange(0, signal24stepsFs / 2, signal24stepsFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 2].set_xlabel('częstotliwość [Hz]')
    axs[3, 2].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('decymacja - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

# interpolacja liniowa
rates = [2000, 4000, 8000, 11999, 16000, 16953, 24000, 41000]
for signal, signalFs, label in zip(signals, signalsFs, signalLabels):
    originalSignal = signal.copy()

    signal2000, signal2000Fs = interpolate(signal, signalFs, 2000 / signalFs, 'linear')
    signal4000, signal4000Fs = interpolate(signal, signalFs, 4000 / signalFs, 'linear')
    signal8000, signal8000Fs = interpolate(signal, signalFs, 8000 / signalFs, 'linear')
    signal11999, signal11999Fs = interpolate(signal, signalFs, 11999 / signalFs, 'linear')
    signal16000, signal16000Fs = interpolate(signal, signalFs, 16000 / signalFs, 'linear')
    signal16953, signal16953Fs = interpolate(signal, signalFs, 16953 / signalFs, 'linear')
    signal24000, signal24000Fs = interpolate(signal, signalFs, 24000 / signalFs, 'linear')
    signal41000, signal41000Fs = interpolate(signal, signalFs, 41000 / signalFs, 'linear')

    fig, axs = plt.subplots(4, 4, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("linear interpolation: {}".format(label))

    axs[0, 0].plot(np.arange(0, signal2000.shape[0]) / signal2000Fs, signal2000, color='black')
    axs[0, 0].set_title('2000')
    axs[0, 0].set_xlim(0, 0.01)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, signal4000.shape[0]) / signal4000Fs, signal4000, color='black')
    axs[0, 1].set_title('4000')
    axs[0, 1].set_xlim(0, 0.01)
    axs[0, 1].set_xlabel('czas trwania [s]')

    axs[0, 2].plot(np.arange(0, signal8000.shape[0]) / signal8000Fs, signal8000, color='black')
    axs[0, 2].set_title('8000')
    axs[0, 2].set_xlim(0, 0.01)
    axs[0, 2].set_xlabel('czas trwania [s]')

    axs[0, 3].plot(np.arange(0, signal11999.shape[0]) / signal11999Fs, signal11999, color='black')
    axs[0, 3].set_title('11999')
    axs[0, 3].set_xlim(0, 0.01)
    axs[0, 3].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(signal2000, fsize)
    axs[1, 0].plot(np.arange(0, signal2000Fs / 2, signal2000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal4000, fsize)
    axs[1, 1].plot(np.arange(0, signal4000Fs / 2, signal4000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal8000, fsize)
    axs[1, 2].plot(np.arange(0, signal8000Fs / 2, signal8000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 2].set_xlabel('częstotliwość [Hz]')
    axs[1, 2].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal11999, fsize)
    axs[1, 3].plot(np.arange(0, signal11999Fs / 2, signal11999Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 3].set_xlabel('częstotliwość [Hz]')
    axs[1, 3].set_ylabel('[dB]')

    axs[2, 0].plot(np.arange(0, signal16000.shape[0]) / signal16000Fs, signal16000, color='black')
    axs[2, 0].set_title('16000')
    axs[2, 0].set_xlim(0, 0.01)
    axs[2, 0].set_xlabel('czas trwania [s]')

    axs[2, 1].plot(np.arange(0, signal16953.shape[0]) / signal16953Fs, signal16953, color='black')
    axs[2, 1].set_title('16953')
    axs[2, 1].set_xlim(0, 0.01)
    axs[2, 1].set_xlabel('czas trwania [s]')

    axs[2, 2].plot(np.arange(0, signal24000.shape[0]) / signal24000Fs, signal24000, color='black')
    axs[2, 2].set_title('24000')
    axs[2, 2].set_xlim(0, 0.01)
    axs[2, 2].set_xlabel('czas trwania [s]')

    axs[2, 3].plot(np.arange(0, signal41000.shape[0]) / signal41000Fs, signal41000, color='black')
    axs[2, 3].set_title('41000')
    axs[2, 3].set_xlim(0, 0.01)
    axs[2, 3].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(signal16000, fsize)
    axs[3, 0].plot(np.arange(0, signal16000Fs / 2, signal16000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 0].set_xlabel('częstotliwość [Hz]')
    axs[3, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal16953, fsize)
    axs[3, 1].plot(np.arange(0, signal16953Fs / 2, signal16953Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 1].set_xlabel('częstotliwość [Hz]')
    axs[3, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal24000, fsize)
    axs[3, 2].plot(np.arange(0, signal24000Fs / 2, signal24000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 2].set_xlabel('częstotliwość [Hz]')
    axs[3, 2].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal41000, fsize)
    axs[3, 3].plot(np.arange(0, signal41000Fs / 2, signal41000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 3].set_xlabel('częstotliwość [Hz]')
    axs[3, 3].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('interpolacja liniowa - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

# interpolacja nieliniowa
for signal, signalFs, label in zip(signals, signalsFs, signalLabels):
    originalSignal = signal.copy()

    signal2000, signal2000Fs = interpolate(signal, signalFs, 2000 / signalFs, 'cubic')
    signal4000, signal4000Fs = interpolate(signal, signalFs, 4000 / signalFs, 'cubic')
    signal8000, signal8000Fs = interpolate(signal, signalFs, 8000 / signalFs, 'cubic')
    signal11999, signal11999Fs = interpolate(signal, signalFs, 11999 / signalFs, 'cubic')
    signal16000, signal16000Fs = interpolate(signal, signalFs, 16000 / signalFs, 'cubic')
    signal16953, signal16953Fs = interpolate(signal, signalFs, 16953 / signalFs, 'cubic')
    signal24000, signal24000Fs = interpolate(signal, signalFs, 24000 / signalFs, 'cubic')
    signal41000, signal41000Fs = interpolate(signal, signalFs, 41000 / signalFs, 'cubic')

    fig, axs = plt.subplots(4, 4, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("nonlinear interpolation: {}".format(label))

    axs[0, 0].plot(np.arange(0, signal2000.shape[0]) / signal2000Fs, signal2000, color='black')
    axs[0, 0].set_title('2000')
    axs[0, 0].set_xlim(0, 0.01)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, signal4000.shape[0]) / signal4000Fs, signal4000, color='black')
    axs[0, 1].set_title('4000')
    axs[0, 1].set_xlim(0, 0.01)
    axs[0, 1].set_xlabel('czas trwania [s]')

    axs[0, 2].plot(np.arange(0, signal8000.shape[0]) / signal8000Fs, signal8000, color='black')
    axs[0, 2].set_title('8000')
    axs[0, 2].set_xlim(0, 0.01)
    axs[0, 2].set_xlabel('czas trwania [s]')

    axs[0, 3].plot(np.arange(0, signal11999.shape[0]) / signal11999Fs, signal11999, color='black')
    axs[0, 3].set_title('11999')
    axs[0, 3].set_xlim(0, 0.01)
    axs[0, 3].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(signal2000, fsize)
    axs[1, 0].plot(np.arange(0, signal2000Fs / 2, signal2000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal4000, fsize)
    axs[1, 1].plot(np.arange(0, signal4000Fs / 2, signal4000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal8000, fsize)
    axs[1, 2].plot(np.arange(0, signal8000Fs / 2, signal8000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 2].set_xlabel('częstotliwość [Hz]')
    axs[1, 2].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal11999, fsize)
    axs[1, 3].plot(np.arange(0, signal11999Fs / 2, signal11999Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 3].set_xlabel('częstotliwość [Hz]')
    axs[1, 3].set_ylabel('[dB]')

    axs[2, 0].plot(np.arange(0, signal16000.shape[0]) / signal16000Fs, signal16000, color='black')
    axs[2, 0].set_title('16000')
    axs[2, 0].set_xlim(0, 0.01)
    axs[2, 0].set_xlabel('czas trwania [s]')

    axs[2, 1].plot(np.arange(0, signal16953.shape[0]) / signal16953Fs, signal16953, color='black')
    axs[2, 1].set_title('16953')
    axs[2, 1].set_xlim(0, 0.01)
    axs[2, 1].set_xlabel('czas trwania [s]')

    axs[2, 2].plot(np.arange(0, signal24000.shape[0]) / signal24000Fs, signal24000, color='black')
    axs[2, 2].set_title('24000')
    axs[2, 2].set_xlim(0, 0.01)
    axs[2, 2].set_xlabel('czas trwania [s]')

    axs[2, 3].plot(np.arange(0, signal41000.shape[0]) / signalFs, signal41000, color='black')
    axs[2, 3].set_title('41000')
    axs[2, 3].set_xlim(0, 0.01)
    axs[2, 3].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(signal16000, fsize)
    axs[3, 0].plot(np.arange(0, signal16000Fs / 2, signal16000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 0].set_xlabel('częstotliwość [Hz]')
    axs[3, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal16953, fsize)
    axs[3, 1].plot(np.arange(0, signal16953Fs / 2, signal16953Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 1].set_xlabel('częstotliwość [Hz]')
    axs[3, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal24000, fsize)
    axs[3, 2].plot(np.arange(0, signal24000Fs / 2, signal24000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 2].set_xlabel('częstotliwość [Hz]')
    axs[3, 2].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(signal41000, fsize)
    axs[3, 3].plot(np.arange(0, signal41000Fs / 2, signal41000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 3].set_xlabel('częstotliwość [Hz]')
    axs[3, 3].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('interpolacja nieliniowa - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

singHigh, singHighFs = sf.read('sing_high2.wav', dtype='float32')
singMedium, singMediumFs = sf.read('sing_medium1.wav', dtype='float32')
singLow, singLowFs = sf.read('sing_low1.wav', dtype='float32')

sings = [singHigh, singMedium, singLow]
singsFs = [singHighFs, singMediumFs, singLowFs]
singLabels = ['high', 'medium', 'low']

# kwantyzacja
for sing, singFs, label in zip(sings, singsFs, singLabels):
    originalSing = sing.copy()

    sing4bits = quantize(originalSing, 4)
    sing8bits = quantize(originalSing, 8)

    if label == 'high':
        timeMargin = [2, 2.01]
    elif label == 'medium':
        timeMargin = [2, 2.01]
    else:
        timeMargin = [2, 3]

    fig, axs = plt.subplots(2, 2, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("quantization: {}".format(label))

    axs[0, 0].plot(np.arange(0, sing4bits.shape[0]) / singFs, sing4bits, color='black')
    axs[0, 0].set_title('4 bits')
    axs[0, 0].set_xlim(timeMargin)
    axs[0, 0].set_xticks(timeMargin)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, sing8bits.shape[0]) / singFs, sing8bits, color='black')
    axs[0, 1].set_title('8 bits')
    axs[0, 1].set_xlim(timeMargin)
    axs[0, 1].set_xticks(timeMargin)
    axs[0, 1].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(sing4bits, fsize)
    axs[1, 0].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing8bits, fsize)
    axs[1, 1].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('kwantyzacja - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

# decymacja
for sing, singFs, label in zip(sings, singsFs, singLabels):
    originalSing = sing.copy()

    sing4steps, sing4stepsFs = decimate(originalSing, singFs, 4)
    sing6steps, sing6stepsFs = decimate(originalSing, singFs, 6)
    sing10steps, sing10stepsFs = decimate(originalSing, singFs, 10)
    sing24steps, sing24stepsFs = decimate(originalSing, singFs, 24)

    timeMargin = [0.1, 0.15]

    fig, axs = plt.subplots(4, 2, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("decimation: {}".format(label))

    axs[0, 0].plot(np.arange(0, sing4steps.shape[0]) / sing4stepsFs, sing4steps, color='black')
    axs[0, 0].set_title('4 steps')
    axs[0, 0].set_xlim(timeMargin)
    axs[0, 0].set_xticks(timeMargin)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, sing6steps.shape[0]) / sing6stepsFs, sing6steps, color='black')
    axs[0, 1].set_title('6 steps')
    axs[0, 1].set_xlim(timeMargin)
    axs[0, 1].set_xticks(timeMargin)
    axs[0, 1].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(sing4steps, fsize)
    axs[1, 0].plot(np.arange(0, sing4stepsFs / 2, sing4stepsFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing6steps, fsize)
    axs[1, 1].plot(np.arange(0, sing6stepsFs / 2, sing6stepsFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    axs[2, 0].plot(np.arange(0, sing10steps.shape[0]) / sing10stepsFs, sing10steps, color='black')
    axs[2, 0].set_title('10 steps')
    axs[2, 0].set_xlim(timeMargin)
    axs[2, 0].set_xticks(timeMargin)
    axs[2, 0].set_xlabel('czas trwania [s]')

    axs[2, 1].plot(np.arange(0, sing24steps.shape[0]) / sing24stepsFs, sing24steps, color='black')
    axs[2, 1].set_title('24 steps')
    axs[2, 1].set_xlim(timeMargin)
    axs[2, 1].set_xticks(timeMargin)
    axs[2, 1].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(sing10steps, fsize)
    axs[3, 0].plot(np.arange(0, sing10stepsFs / 2, sing10stepsFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 0].set_xlabel('częstotliwość [Hz]')
    axs[3, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing24steps, fsize)
    axs[3, 1].plot(np.arange(0, sing24stepsFs / 2, sing24stepsFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 1].set_xlabel('częstotliwość [Hz]')
    axs[3, 1].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('decymacja - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

# interpolacja liniowa
for sing, singFs, label in zip(sings, singsFs, singLabels):
    originalSing = sing.copy()

    sing4000, sing4000Fs = interpolate(originalSing, singFs, 4000 / singFs, 'linear')
    sing8000, sing8000Fs = interpolate(originalSing, singFs, 8000 / singFs, 'linear')
    sing11999, sing11999Fs = interpolate(originalSing, singFs, 11999 / singFs, 'linear')
    sing16000, sing16000Fs = interpolate(originalSing, singFs, 16000 / singFs, 'linear')
    sing16953, sing16953Fs = interpolate(originalSing, singFs, 16953 / singFs, 'linear')

    timeMargin = [0.1, 0.125]

    fig, axs = plt.subplots(4, 3, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("linear interpolation: {}".format(label))

    axs[0, 0].plot(np.arange(0, originalSing.shape[0]) / singFs, originalSing, color='black')
    axs[0, 0].set_title('original')
    axs[0, 0].set_xlim(timeMargin)
    axs[0, 0].set_xticks(timeMargin)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, sing4000.shape[0]) / sing4000Fs, sing4000, color='black')
    axs[0, 1].set_title('4000')
    axs[0, 1].set_xlim(timeMargin)
    axs[0, 1].set_xticks(timeMargin)
    axs[0, 1].set_xlabel('czas trwania [s]')

    axs[0, 2].plot(np.arange(0, sing8000.shape[0]) / sing8000Fs, sing8000, color='black')
    axs[0, 2].set_title('8000')
    axs[0, 2].set_xlim(timeMargin)
    axs[0, 2].set_xticks(timeMargin)
    axs[0, 2].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(originalSing, fsize)
    axs[1, 0].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing4000, fsize)
    axs[1, 1].plot(np.arange(0, sing4000Fs / 2, sing4000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing8000, fsize)
    axs[1, 2].plot(np.arange(0, sing8000Fs / 2, sing8000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 2].set_xlabel('częstotliwość [Hz]')
    axs[1, 2].set_ylabel('[dB]')

    axs[2, 0].plot(np.arange(0, sing11999.shape[0]) / sing11999Fs, sing11999, color='black')
    axs[2, 0].set_title('11999')
    axs[2, 0].set_xlim(timeMargin)
    axs[2, 0].set_xticks(timeMargin)
    axs[2, 0].set_xlabel('czas trwania [s]')

    axs[2, 1].plot(np.arange(0, sing16000.shape[0]) / sing16000Fs, sing16000, color='black')
    axs[2, 1].set_title('16000')
    axs[2, 1].set_xlim(timeMargin)
    axs[2, 1].set_xticks(timeMargin)
    axs[2, 1].set_xlabel('czas trwania [s]')

    axs[2, 2].plot(np.arange(0, sing16953.shape[0]) / sing16953Fs, sing16953, color='black')
    axs[2, 2].set_title('16953')
    axs[2, 2].set_xlim(timeMargin)
    axs[2, 2].set_xticks(timeMargin)
    axs[2, 2].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(sing11999, fsize)
    axs[3, 0].plot(np.arange(0, sing11999Fs / 2, sing11999Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 0].set_xlabel('częstotliwość [Hz]')
    axs[3, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing16000, fsize)
    axs[3, 1].plot(np.arange(0, sing16000Fs / 2, sing16000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 1].set_xlabel('częstotliwość [Hz]')
    axs[3, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing16953, fsize)
    axs[3, 2].plot(np.arange(0, sing16953Fs / 2, sing16953Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 2].set_xlabel('częstotliwość [Hz]')
    axs[3, 2].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('interpolacja liniowa - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

# interpolacja nieliniowa
for sing, singFs, label in zip(sings, singsFs, singLabels):
    originalSing = sing.copy()

    sing4000, sing4000Fs = interpolate(originalSing, singFs, 4000 / singFs, 'cubic')
    sing8000, sing8000Fs = interpolate(originalSing, singFs, 8000 / singFs, 'cubic')
    sing11999, sing11999Fs = interpolate(originalSing, singFs, 11999 / singFs, 'cubic')
    sing16000, sing16000Fs = interpolate(originalSing, singFs, 16000 / singFs, 'cubic')
    sing16953, sing16953Fs = interpolate(originalSing, singFs, 16953 / singFs, 'cubic')

    timeMargin = [0.1, 0.125]

    fig, axs = plt.subplots(4, 3, figsize=(20, 15))
    fig.tight_layout(pad=5.0)
    fig.suptitle("nonlinear interpolation: {}".format(label))

    axs[0, 0].plot(np.arange(0, originalSing.shape[0]) / singFs, originalSing, color='black')
    axs[0, 0].set_title('original')
    axs[0, 0].set_xlim(timeMargin)
    axs[0, 0].set_xticks(timeMargin)
    axs[0, 0].set_xlabel('czas trwania [s]')

    axs[0, 1].plot(np.arange(0, sing4000.shape[0]) / sing4000Fs, sing4000, color='black')
    axs[0, 1].set_title('4000')
    axs[0, 1].set_xlim(timeMargin)
    axs[0, 1].set_xticks(timeMargin)
    axs[0, 1].set_xlabel('czas trwania [s]')

    axs[0, 2].plot(np.arange(0, sing8000.shape[0]) / sing8000Fs, sing8000, color='black')
    axs[0, 2].set_title('8000')
    axs[0, 2].set_xlim(timeMargin)
    axs[0, 2].set_xticks(timeMargin)
    axs[0, 2].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(originalSing, fsize)
    axs[1, 0].plot(np.arange(0, singFs / 2, singFs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])), color='red')
    axs[1, 0].set_xlabel('częstotliwość [Hz]')
    axs[1, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing4000, fsize)
    axs[1, 1].plot(np.arange(0, sing4000Fs / 2, sing4000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 1].set_xlabel('częstotliwość [Hz]')
    axs[1, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing8000, fsize)
    axs[1, 2].plot(np.arange(0, sing8000Fs / 2, sing8000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[1, 2].set_xlabel('częstotliwość [Hz]')
    axs[1, 2].set_ylabel('[dB]')

    axs[2, 0].plot(np.arange(0, sing11999.shape[0]) / sing11999Fs, sing11999, color='black')
    axs[2, 0].set_title('11999')
    axs[2, 0].set_xlim(timeMargin)
    axs[2, 0].set_xticks(timeMargin)
    axs[2, 0].set_xlabel('czas trwania [s]')

    axs[2, 1].plot(np.arange(0, sing16000.shape[0]) / sing16000Fs, sing16000, color='black')
    axs[2, 1].set_title('16000')
    axs[2, 1].set_xlim(timeMargin)
    axs[2, 1].set_xticks(timeMargin)
    axs[2, 1].set_xlabel('czas trwania [s]')

    axs[2, 2].plot(np.arange(0, sing16953.shape[0]) / sing16953Fs, sing16953, color='black')
    axs[2, 2].set_title('16953')
    axs[2, 2].set_xlim(timeMargin)
    axs[2, 2].set_xticks(timeMargin)
    axs[2, 2].set_xlabel('czas trwania [s]')

    yf = scipy.fftpack.fft(sing11999, fsize)
    axs[3, 0].plot(np.arange(0, sing11999Fs / 2, sing11999Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 0].set_xlabel('częstotliwość [Hz]')
    axs[3, 0].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing16000, fsize)
    axs[3, 1].plot(np.arange(0, sing16000Fs / 2, sing16000Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 1].set_xlabel('częstotliwość [Hz]')
    axs[3, 1].set_ylabel('[dB]')

    yf = scipy.fftpack.fft(sing16953, fsize)
    axs[3, 2].plot(np.arange(0, sing16953Fs / 2, sing16953Fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])),
                   color='red')
    axs[3, 2].set_xlabel('częstotliwość [Hz]')
    axs[3, 2].set_ylabel('[dB]')

    memfile = BytesIO()
    plt.savefig(memfile, bbox_inches='tight')
    document.add_heading('interpolacja nieliniowa - {}'.format(label), 2)
    document.add_picture(memfile, width=Inches(6))
    plt.close(fig)

document.save('lab05.docx')
