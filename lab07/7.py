from docx import Document
from docx.shared import Inches
from io import BytesIO
import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import sounddevice as sd
import soundfile as sf


def quantize(signal, targetBits):
    targetMax = 2 ** targetBits - 1
    quantizedSignal = np.round(signal * (targetMax / 2)) / (targetMax / 2)
    return quantizedSignal


def encodeALaw(x, quantizationLevel, A=87.6):
    signs = np.sign(x)
    compression = np.log(1 + A * np.abs(x)) / np.log(1 + A)
    quantized = np.round(compression * (2 ** quantizationLevel - 1))
    return signs * quantized


def decodeALaw(y, quantizationLevel, A=87.6):
    signs = np.sign(y)
    expansion = (np.exp(np.abs(y) / (2 ** quantizationLevel - 1) * np.log(1 + A)) - 1) / A
    return signs * expansion


def encodeMuLaw(x, quantizationLevel, mu=255):
    signs = np.sign(x)
    compression = np.log1p(mu * np.abs(x)) / np.log1p(mu)
    quantized = np.round(compression * (2 ** quantizationLevel - 1))
    return signs * quantized


def decodeMuLaw(y, quantizationLevel, mu=255):
    signs = np.sign(y)
    expansion = (np.expm1(np.abs(y) / (2 ** quantizationLevel - 1) * np.log1p(mu))) / mu
    return signs * expansion


def encodeDPCM(x, quantizationLevel):
    y = np.zeros(x.shape)
    e = 0
    for i in range(x.shape[0]):
        y[i] = quantize(x[i] - e, quantizationLevel)
        e += y[i]
    return y


def decodeDPCM(y):
    x = np.zeros(y.shape)
    x[-1] = 0
    for i in range(y.shape[0]):
        x[i] = y[i] + x[i - 1]
    return x


def predictor(x):
    if len(x) == 0:
        return 0
    return np.median(x)


def encodeDPCMpred(x, quantizationLevel, n=30):
    y = np.zeros(x.shape)
    xp = np.zeros(x.shape)
    e = 0
    for i in range(1, x.shape[0]):
        y[i] = quantize(x[i] - e, quantizationLevel)
        xp[i] = y[i] + e
        idx = (np.arange(i - n, i, 1, dtype=int) + 1)
        idx = np.delete(idx, idx < 0)
        e = predictor(xp[idx])
    return y


def decodeDCPMpred(y, n=30):
    x = np.zeros(y.shape)
    xp = np.zeros(y.shape)
    e = 0
    for i in range(1, y.shape[0]):
        xp[i] = y[i] + e
        x[i] = xp[i]
        idx = np.arange(i - n, i, 1, dtype=int) + 1
        idx = np.delete(idx, idx < 0)
        e = predictor(xp[idx])
    return x


document = Document()
document.add_heading('Laboratorium 07', 0)

singHigh, singHighFs = sf.read('sing_high2.wav', dtype='float32')
singMedium, singMediumFs = sf.read('sing_medium1.wav', dtype='float32')
singLow, singLowFs = sf.read('sing_low1.wav', dtype='float32')

sings = [singHigh, singMedium, singLow]
singsNames = ['singHigh', 'singMedium', 'singLow']
singsFs = [singHighFs, singMediumFs, singLowFs]
quantizationLevels = [8, 7, 6, 5, 4, 3, 2]
i = [0, 1, 2, 3, 4, 5, 6]
for sing, singName, singFs in zip(sings, singsNames, singsFs):
    document.add_heading('{} - A-law, mu-law, DPCM bez predykcji, DPCM z predykcją'.format(singName), 2)
    for quantizationLevel in quantizationLevels:
        originalSing = sing.copy()

        compressedALaw = encodeALaw(originalSing, quantizationLevel)
        decompressedALaw = decodeALaw(compressedALaw, quantizationLevel)
        compressedMuLaw = encodeMuLaw(originalSing, quantizationLevel)
        decompressedMuLaw = decodeMuLaw(compressedMuLaw, quantizationLevel)
        compressedDPCM = encodeDPCM(originalSing, quantizationLevel)
        decompressedDPCM = decodeDPCM(compressedDPCM)
        compressedDPCMpred = encodeDPCMpred(originalSing, quantizationLevel)
        decompressedDPCMpred = decodeDCPMpred(compressedDPCMpred)

        memfile = BytesIO()
        plt.plot(np.arange(0, originalSing.shape[0]) / singFs, originalSing, color='blue')
        plt.plot(np.arange(0, decompressedALaw.shape[0]) / singFs, decompressedALaw, color='red')
        plt.plot(np.arange(0, decompressedMuLaw.shape[0]) / singFs, decompressedMuLaw, color='green')
        plt.plot(np.arange(0, decompressedDPCM.shape[0]) / singFs, decompressedDPCM, color='purple')
        plt.plot(np.arange(0, decompressedDPCMpred.shape[0]) / singFs, decompressedDPCMpred, color='black')
        plt.xlim(1, 1.01)
        plt.title('{}-bit quantization'.format(quantizationLevel))
        plt.legend(['original', 'A-law', 'mu-law', 'DPCM', 'DPCM pred'])
        plt.savefig(memfile, bbox_inches='tight', dpi=500)
        document.add_picture(memfile, width=Inches(3.5))
        plt.close()

document.save('lab07.docx')
