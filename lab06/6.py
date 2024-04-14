from docx import Document
from docx.shared import Inches
from io import BytesIO
import numpy as np
import matplotlib.pyplot as plt


def calculateCompressionRatio(img, imgEncoded):
    originalSize = img.size
    encodedSize = len(imgEncoded)

    compressionRatio = abs(originalSize) / abs(encodedSize)
    compressionPercentage = 100 * (abs(encodedSize) / abs(originalSize))

    return compressionRatio, compressionPercentage


def encodeRLE(img):
    pixels = img.flatten()
    encoded = []
    count = 1
    prevPixel = pixels[0]

    for pixel in pixels[1:]:
        if np.array_equal(pixel, prevPixel):
            count += 1
        else:
            encoded.append((count, prevPixel))
            prevPixel = pixel
            count = 1
    encoded.append((count, prevPixel))
    return encoded


def decodeRLE(encodedPixels, shape):
    decodedImage = []

    for count, pixel in encodedPixels:
        decodedImage.extend([pixel] * count)
    return np.array(decodedImage).reshape(shape)


def encodeByteRun(img):
    pixels = img.flatten()
    encoded = []
    count = 1
    prevPixel = pixels[0]

    for pixel in pixels[1:]:
        if pixel == prevPixel and count < 255:
            count += 1
        else:
            encoded.append(count)
            encoded.append(prevPixel)
            prevPixel = pixel
            count = 1
    encoded.append(count)
    encoded.append(prevPixel)
    return np.array(encoded, dtype=np.uint8)


def decodeByteRun(imgEncoded):
    pixels = []
    it = iter(imgEncoded)
    for count, pixel in zip(it, it):
        pixels.extend([pixel] * count)
    return np.array(pixels, dtype=np.uint8)


document = Document()
document.add_heading('Laboratorium 06', 0)

images = ['technical.jpg', 'colored.jpg', 'document.jpg']
for image in images:
    img = plt.imread(image)
    img = img.astype(np.uint8)
    document.add_heading('{} - RLE i ByteRun'.format(image), 2)

    memfile = BytesIO()
    plt.figure(figsize=(10, 8))
    plt.imshow(img)
    plt.title('Original')
    plt.axis('off')
    plt.savefig(memfile, bbox_inches='tight', dpi=500)
    document.add_picture(memfile, width=Inches(3.5))
    document.add_paragraph('Wymiary oryginału: {}'.format(img.shape))
    plt.close()

    imgEncodedRLE = encodeRLE(img.copy())
    compressionRatio, compressionPercentage = calculateCompressionRatio(img, imgEncodedRLE)
    print(f'Stopień kompresji RLE: {compressionRatio:.4f} ({compressionPercentage:.2f}%)')
    imgDecodedRLE = decodeRLE(imgEncodedRLE, img.shape)

    memfile = BytesIO()
    plt.figure(figsize=(10, 8))
    plt.imshow(imgDecodedRLE)
    plt.title('Decoded RLE')
    plt.axis('off')
    plt.text(30, 50, f'Stopień kompresji RLE: {compressionRatio:.4f} ({compressionPercentage:.2f}%)',
             fontsize=14, bbox=dict(facecolor='red', alpha=0.5))
    plt.savefig(memfile, bbox_inches='tight', dpi=500)
    document.add_picture(memfile, width=Inches(3.5))
    document.add_paragraph('Wymiary dekompresji RLE: {}'.format(imgDecodedRLE.shape))
    plt.close()

    imgEncodedByteRun = encodeByteRun(img.copy())
    compressionRatio, compressionPercentage = calculateCompressionRatio(img, imgEncodedByteRun)
    print(f'Stopień kompresji ByteRun: {compressionRatio:.4f} ({compressionPercentage:.2f}%)')
    imgDecodedByteRun = decodeByteRun(imgEncodedByteRun).reshape(img.shape)

    memfile = BytesIO()
    plt.figure(figsize=(10, 8))
    plt.imshow(imgDecodedByteRun)
    plt.title('Decoded ByteRun')
    plt.axis('off')
    plt.text(30, 50, f'Stopień kompresji ByteRun: {compressionRatio:.4f} ({compressionPercentage:.2f}%)',
             fontsize=14, bbox=dict(facecolor='red', alpha=0.5))
    plt.savefig(memfile, bbox_inches='tight', dpi=500)
    document.add_picture(memfile, width=Inches(3.5))
    document.add_paragraph('Wymiary dekompresji RLE: {}'.format(imgDecodedRLE.shape))
    plt.close()

document.save('lab06.docx')
