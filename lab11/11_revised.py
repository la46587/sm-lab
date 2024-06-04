import cv2
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from io import BytesIO


def calcMSE(imgOriginal, imgModified):
    return np.mean((imgOriginal.astype("float") - imgModified.astype("float")) ** 2)


def water_mark(img, mask, alpha=0.25):
    assert (img.shape[0] == mask.shape[0]) and (img.shape[1] == mask.shape[1]), "Wrong size"
    if len(img.shape) < 3:
        flag = True
        t_img = cv2.cvtColor(img, cv2.COLOR_GRAY2RGBA)
    else:
        flag = False
        t_img = cv2.cvtColor(img, cv2.COLOR_RGB2RGBA)
    if mask.dtype == bool:
        t_mask = cv2.cvtColor((mask * 255).astype(np.uint8), cv2.COLOR_GRAY2RGBA)
    elif mask.dtype == np.uint8:
        if len(mask.shape) < 3:
            t_mask = cv2.cvtColor(mask.astype(np.uint8), cv2.COLOR_GRAY2RGBA)
        else:
            t_mask = cv2.cvtColor(mask.astype(np.uint8), cv2.COLOR_RGB2RGBA)
    else:
        if len(mask.shape) < 3:
            t_mask = cv2.cvtColor((mask * 255).astype(np.uint8), cv2.COLOR_GRAY2RGBA)
        else:
            t_mask = cv2.cvtColor((mask * 255).astype(np.uint8), cv2.COLOR_RGB2RGBA)
    t_out = cv2.addWeighted(t_img, 1, t_mask, alpha, 0)
    if flag:
        out = cv2.cvtColor(t_out, cv2.COLOR_RGBA2GRAY)
    else:
        out = cv2.cvtColor(t_out, cv2.COLOR_RGBA2RGB)
    return out


def put_data(img, data, binary_mask=np.uint8(1)):
    assert img.dtype == np.uint8, "img wrong data type"
    assert binary_mask.dtype == np.uint8, "binary_mask wrong data type"
    un_binary_mask = np.unpackbits(binary_mask)
    if data.dtype != bool:
        unpacked_data = np.unpackbits(data)
    else:
        unpacked_data = data
    dataspace = img.shape[0]*img.shape[1]*np.sum(un_binary_mask)
    assert (dataspace >= unpacked_data.size), "too much data"
    if dataspace == unpacked_data.size:
        prepared_data = unpacked_data.reshape(img.shape[0], img.shape[1], np.sum(un_binary_mask)).astype(np.uint8)
    else:
        prepared_data = np.resize(unpacked_data, (img.shape[0], img.shape[1], np.sum(un_binary_mask))).astype(np.uint8)
    mask = np.full((img.shape[0], img.shape[1]), binary_mask)
    img = np.bitwise_and(img, np.invert(mask))
    bv = 0
    for i, b in enumerate(un_binary_mask[::-1]):
        if b:
            temp = prepared_data[:, :, bv]
            temp = np.left_shift(temp, i)
            img = np.bitwise_or(img, temp)
            bv += 1
    return img


def pop_data(img, binary_mask=np.uint8(1), out_shape=None):
    un_binary_mask = np.unpackbits(binary_mask)
    data = np.zeros((img.shape[0], img.shape[1], np.sum(un_binary_mask))).astype(np.uint8)
    bv = 0
    for i, b in enumerate(un_binary_mask[::-1]):
        if b:
            mask = np.full((img.shape[0], img.shape[1]), 2 ** i)
            temp = np.bitwise_and(img, mask)
            data[:, :, bv] = temp[:, :].astype(np.uint8)
            bv += 1
    if out_shape is not None:
        tmp = np.packbits(data.flatten())
        tmp = tmp[:np.prod(out_shape)]
        data = tmp.reshape(out_shape)
    return data


document = Document()
document.add_heading('Laboratorium 11', 0)

img = cv2.imread('zut.jpg')
red_img = img[:, :, 2]

secret = cv2.imread('logo.jpg')
secret_bits = np.unpackbits(secret).astype(bool)
binary_mask = np.array([1], dtype=np.uint8)
encoded_red_img = put_data(red_img, secret_bits, binary_mask)
encoded_img = img.copy()
encoded_img[:, :, 2] = encoded_red_img
decoded_img = pop_data(encoded_red_img, binary_mask, out_shape=secret.shape)

memfile = BytesIO()
fig, axs = plt.subplots(1, 3)
fig.tight_layout()
axs[0].imshow(cv2.cvtColor(img, cv2.COLOR_BGR2RGB))
axs[0].axis('off')
axs[0].set_title('Original')
axs[1].imshow(cv2.cvtColor(encoded_img, cv2.COLOR_BGR2RGB))
axs[1].axis('off')
axs[1].set_title('Encoded')
axs[2].imshow(cv2.cvtColor(decoded_img, cv2.COLOR_BGR2RGB))
axs[2].axis('off')
axs[2].set_title('Decoded information')
plt.savefig(memfile)
document.add_heading('Kodowanie obrazu binarnego', 1)
document.add_picture(memfile)
plt.close(fig)

img = cv2.imread('zut.jpg')
blue_img = img[:, :, 0]
green_img = img[:, :, 1]
red_img = img[:, :, 2]

secret = cv2.imread('logo_colored.jpg')
blue_secret = secret[:, :, 0]
green_secret = secret[:, :, 1]
red_secret = secret[:, :, 2]

blue_secret_bits = np.unpackbits(blue_secret).astype(bool)
green_secret_bits = np.unpackbits(green_secret).astype(bool)
red_secret_bits = np.unpackbits(red_secret).astype(bool)
blue_binary_mask = np.array([1], dtype=np.uint8)
green_binary_mask = np.array([3], dtype=np.uint8)
red_binary_mask = np.array([5], dtype=np.uint8)
encoded_blue_img = put_data(blue_img, blue_secret_bits, blue_binary_mask)
encoded_green_img = put_data(green_img, green_secret_bits, green_binary_mask)
encoded_red_img = put_data(red_img, red_secret_bits, red_binary_mask)

encoded_img = img.copy()
encoded_img[:, :, 0] = encoded_blue_img
encoded_img[:, :, 1] = encoded_green_img
encoded_img[:, :, 2] = encoded_red_img

decoded_blue = pop_data(encoded_img[:, :, 0], blue_binary_mask, out_shape=(secret.shape[0], secret.shape[1], secret.shape[2]//3))
decoded_green = pop_data(encoded_img[:, :, 1], green_binary_mask, out_shape=(secret.shape[0], secret.shape[1], secret.shape[2]//3))
decoded_red = pop_data(encoded_img[:, :, 2], red_binary_mask, out_shape=(secret.shape[0], secret.shape[1], secret.shape[2]//3))

decoded_img = np.concatenate((decoded_blue, decoded_green, decoded_red), axis=2)

memfile = BytesIO()
fig, axs = plt.subplots(1, 3)
fig.tight_layout()
axs[0].imshow(cv2.cvtColor(img, cv2.COLOR_BGR2RGBA))
axs[0].axis('off')
axs[0].set_title('Original')
axs[1].imshow(cv2.cvtColor(encoded_img, cv2.COLOR_BGRA2RGBA))
axs[1].axis('off')
axs[1].set_title('Encoded 3 channels')
axs[2].imshow(cv2.cvtColor(decoded_img, cv2.COLOR_BGRA2RGBA))
axs[2].axis('off')
axs[2].set_title('Decoded image')
plt.savefig(memfile)
document.add_heading('Kodowanie obrazu kolorowego', 1)
document.add_picture(memfile)
plt.close(fig)

img = cv2.imread('opel.jpeg')
mask = cv2.imread('watermark.png')
imgWatermarkLow = water_mark(img, mask, alpha=0.1)
imgWatermarkMedium = water_mark(img, mask, alpha=0.25)
imgWatermarkHigh = water_mark(img, mask, alpha=0.5)

memfile = BytesIO()
fig, axs = plt.subplots(2, 2)
axs[0, 0].imshow(cv2.cvtColor(img, cv2.COLOR_BGR2RGB))
axs[0, 0].axis('off')
axs[0, 0].set_title('Original')
axs[0, 1].imshow(cv2.cvtColor(imgWatermarkLow, cv2.COLOR_BGR2RGB))
axs[0, 1].axis('off')
axs[0, 1].set_title('alpha = 0.1')
axs[1, 0].imshow(cv2.cvtColor(imgWatermarkMedium, cv2.COLOR_BGR2RGB))
axs[1, 0].axis('off')
axs[1, 0].set_title('alpha = 0.25')
axs[1, 1].imshow(cv2.cvtColor(imgWatermarkHigh, cv2.COLOR_BGR2RGB))
axs[1, 1].axis('off')
axs[1, 1].set_title('alpha = 0.5')
plt.savefig(memfile)
document.add_heading('Znak wodny')
document.add_picture(memfile)
plt.close(fig)

mseLow = calcMSE(img, imgWatermarkLow)
mseMedium = calcMSE(img, imgWatermarkMedium)
mseHigh = calcMSE(img, imgWatermarkHigh)

document.add_paragraph(f'MSE (alfa 0.1) - {mseLow:.2f}')
document.add_paragraph(f'MSE (alfa 0.25) - {mseMedium:.2f}')
document.add_paragraph(f'MSE (alfa 0.5) - {mseHigh:.2f}')

document.save('lab11bis.docx')
