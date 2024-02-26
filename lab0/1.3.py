from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
import scipy.fftpack
import sounddevice as sd
import soundfile as sf
from io import BytesIO


def plotAudio(Signal, Fs, TimeMargin=[0, 0.02]):
    fsize = 2 ** 8
    plt.figure()
    plt.subplot(3, 1, 1)
    plt.xlim(TimeMargin)
    plt.plot(np.arange(0, data.shape[0]) / Fs, Signal)
    plt.xlabel('czas trwania [s]')
    plt.subplot(2, 1, 2)
    yf = scipy.fftpack.fft(Signal, fsize)
    plt.plot(np.arange(0, fs / 2, fs / fsize), 20 * np.log10(np.abs(yf[:fsize // 2])))
    plt.xlabel('częstotliwość [Hz]')
    plt.ylabel('[dB]')
    plt.show()

document = Document()
document.add_heading('Zmień ten tytuł', 0)  # tworzenie nagłówków, druga wartość to poziom nagłówka

files = ['sin60Hz.wav', 'sin440Hz.wav', 'sin8000Hz.wav']
Margins = [[0, 0.02], [0.133, 0.155]]
for file in files:
    document.add_heading('Plik - {}'.format(file), 2)
    for i, Margin in enumerate(Margins):
        document.add_heading('Time margin {}'.format(Margin), 3)  # nagłówek sekcji, mozę być poziom wyżej
        fig, axs = plt.subplots(2, 1, figsize=(10, 7))  # tworzenie plota

        ############################################################
        # tu wykonujesz jakieś funkcje i rysujesz wykresy
        ############################################################

        fig.suptitle('Time margin {}'.format(Margin))  # tytuł wykresu
        fig.tight_layout(pad=1.5)  # poprawa czytelności
        memfile = BytesIO()  # tworzenie bufora
        fig.savefig(memfile)  # z zapis do bufora

        document.add_picture(memfile, width=Inches(6))  # dodanie obrazu z bufora do pliku

        memfile.close()
        ############################################################
        # tu dodajesz dane tekstowe - wartosci, wyjście funkcji etc.
        document.add_paragraph('wartość losowa = {}'.format(np.random.rand(1)))
        ############################################################

document.save('report.docx')  # zapis do pliku
